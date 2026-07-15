# -*- coding: utf-8 -*-
"""
依据《计算机综合项目实践实验报告模板.docx》的格式，生成 EmailClaw 项目实验报告。
保留模板的封面、郑重声明与全部样式，仅替换正文内容。
"""
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEMPLATE = "docs/reports/计算机综合项目实践实验报告模板.docx"
OUTPUT = "docs/reports/EmailClaw_实验报告.docx"

doc = docx.Document(TEMPLATE)
body = doc.element.body

# ---------------- 字体辅助 ----------------
def set_run_font(run, latin="Times New Roman", ea="宋体", size=10.5, bold=False, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), ea)
    rfonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color

# ---------------- 段落构造 ----------------
def h1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, ea="黑体", size=22, bold=True)
    return p

def h2(text):
    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run(text)
    set_run_font(r, ea="黑体", size=16, bold=True)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, ea="黑体", latin="Times New Roman", size=12, bold=True)
    return p

def para(text, size=10.5, bold=False, indent=True, align="both"):
    p = doc.add_paragraph()
    if align == "both":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)  # 2字符
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Pt(21 + level * 14)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("• " + text)
    set_run_font(r, size=10.5)
    return p

def code_block(text):
    """以等宽字体逐行输出代码/示意图，带浅灰底纹。"""
    lines = text.split("\n")
    # 去掉首尾空行
    while lines and lines[0].strip() == "":
        lines.pop(0)
    while lines and lines[-1].strip() == "":
        lines.pop()
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Pt(14)
        # 底纹
        ppr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        ppr.append(shd)
        r = p.add_run(line if line != "" else " ")
        set_run_font(r, latin="Consolas", ea="宋体", size=9)
    return

def set_cell_text(cell, text, bold=False, size=10, align="left", header=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold or header)
    if header:
        # 表头底纹
        tcpr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E2F3")
        tcpr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tblPr.append(borders)

def add_table(headers, rows, widths=None, size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_text(hdr[i], htext, header=True, align="center", size=size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=size)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Pt(w)
    # 表格后空一行
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    return table

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, ea="楷体", size=9, bold=False, color=RGBColor(0x40, 0x40, 0x40))

def page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

# ====================================================================
# 1) 修改封面
# ====================================================================
def set_para_text_keepfmt(p, new_text):
    """替换段落文字，保留首个 run 的格式。"""
    if not p.runs:
        r = p.add_run(new_text)
        set_run_font(r, ea="宋体", size=15)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""

paras = doc.paragraphs
# [1] 报告标题
set_para_text_keepfmt(paras[1], "智能邮件管家系统实验报告")
# [8] 课程名称
set_para_text_keepfmt(paras[8], "课 程 名 称   ：计算机综合项目实践")
# [11] 学生姓名
set_para_text_keepfmt(paras[11], "学 生 姓 名   ：吕玥  董一鸣")
# [17] 日期
set_para_text_keepfmt(paras[17], "二○二六年六月")
# 声明日期 [42]
set_para_text_keepfmt(paras[42], "本人签名：                     日期：2026年6月6日")

# ====================================================================
# 2) 删除模板正文（从“摘要”标题，index 46 起到结尾的所有段落）
# ====================================================================
# 重新取一次最新 paragraph 列表
all_p = doc.paragraphs
for p in all_p[46:]:
    p._element.getparent().remove(p._element)

print("封面与样式已就绪，开始写入正文……")

# ====================================================================
# 摘要
# ====================================================================
page_break()
h1("摘  要")

para("在电子邮件已成为日常工作与学习核心沟通渠道的今天，邮件过载逐渐成为困扰职场人士与学生群体的普遍问题：每天数十乃至上百封邮件中，真正重要的内容常常被各类营销、通知和垃圾邮件淹没。传统邮件客户端只提供收发与文件夹归档功能，缺乏对邮件内容的语义理解、优先级判断与主动提醒能力，用户往往需要逐封翻阅才能完成筛选与处理，效率低下。")

para("针对这一痛点，本项目设计并实现了一套面向多用户的智能邮件管家系统 EmailClaw。系统通过 IMAP 协议实时监听用户邮箱，自动抓取并解析新邮件；随后依次经过用户自定义规则引擎与智能分析 Agent，完成邮件分类、重要性评分与摘要生成；分析结果以交互式卡片的形式主动推送到飞书机器人，用户可在飞书中直接完成标记已读、标为重点、归档、删除、分类纠错、查看详情与重新分析等操作，所有操作会经后端同步回真实邮箱（IMAP）与数据库，形成完整的“感知—分析—展示—反馈—执行”闭环。智能分析层既内置不依赖大模型的本地关键词规则 Agent，又可选接入 OpenClaw 框架并由其调用 DeepSeek 大模型增强分析质量，大模型调用失败时自动回退本地规则，保证系统的健壮性。")

para("在工程实现上，系统采用 Node.js + TypeScript 技术栈，后端基于 Express 5 提供 REST API，使用 Prisma ORM 操作 PostgreSQL 数据库，邮件收发依托 imap 与 mailparser，飞书交互依托飞书开放平台官方 Node SDK。系统实现了完整的多用户体系：用户密码经 bcrypt 哈希、IMAP 授权码经 AES-256-GCM 加密入库，并基于 JWT 与共享密钥构建了 HTTP、业务、数据三层防御性鉴权，确保多账号场景下的数据隔离与防越权。经端到端联调验证，系统已稳定打通邮件抓取、解析、入库、分析、推送、交互与真实邮箱操作全链路，覆盖了需求规格说明书中 FR-1 至 FR-6 的核心功能。")

para("关键词：智能邮件管家；IMAP；智能分析 Agent；邮件分类；飞书机器人；OpenClaw；多用户鉴权", bold=False, indent=False)

# ====================================================================
# 第1章 概述
# ====================================================================
page_break()
h1("第1章 概述")

h2("1.1 选题")
para("电子邮件诞生至今已成为最重要的异步通信工具之一，被广泛应用于工作协作、商务往来、账号通知、订单物流与营销推广等场景。然而，邮件数量的爆炸式增长带来了严重的信息过载问题。对于一名普通职场人士，每天收到 50～200 封邮件并不罕见，其中混杂着会议通知、项目进展、广告促销、系统通知和垃圾邮件等多种类型。在这样的环境下，用户面临三个突出困难：一是难以从海量邮件中快速识别真正重要、需要立即处理的邮件；二是逐封阅读长篇邮件正文耗时费力，难以快速把握要点；三是传统邮件客户端缺乏个性化的自动分类与处理能力，所有筛选工作都要手工完成。")

para("基于上述背景，本项目选定“基于 Agent 与大语言模型的智能邮件分类与管理系统”作为课题，命名为 EmailClaw（智能邮件管家）。系统的核心目标是：在不改变用户原有邮箱的前提下，作为一个智能中间层实时监听邮箱，借助规则与 AI 能力对每封新邮件自动完成分类、重要性评分与摘要，并将重要邮件以美观的交互式卡片主动推送给用户；用户无需打开邮箱，即可在常用的即时通讯工具中完成对邮件的常见处理操作，且这些操作会真实同步回邮箱服务器。系统进一步支持多用户、用户自定义规则、邮件检索与防御性鉴权，力求贴近真实可用的产品形态。")

para("本课题的意义在于：其一，它是一个完整的、贯通前后端与第三方平台的综合性工程实践，涉及网络协议（IMAP/MIME）、数据库建模、异步并发、第三方开放平台对接、AI Agent 编排与系统安全等多方面知识，能够较好地综合训练软件工程能力；其二，它面向真实痛点，具有实用价值，所采用的“规则引擎 + 本地 Agent + 可选大模型增强 + 自动回退”的分层智能策略，在成本、稳定性与效果之间取得了务实的平衡，对同类 AI 应用的工程化落地具有参考价值。")

h2("1.2 分组及分工")
para("本项目由两名同学协作完成。整体采用前后端分离、各司其职又紧密联调的方式推进，分工如下表所示。下文为聚焦实验内容本身，不再反复强调个人，统一以系统模块视角进行阐述。")
add_table(
    ["成员", "主要职责", "具体工作"],
    [
        ["吕玥", "后端服务与 OpenClaw 接入",
         "负责后端整体架构与核心服务实现，包括多用户 IMAP 连接管理、邮件处理流水线、规则引擎、数据库与 REST API、三层防御性鉴权，以及 OpenClaw + DeepSeek 大模型分析能力的接入、子进程适配与回退机制。"],
        ["董一鸣", "前端服务与模型参数调优",
         "负责飞书机器人前端服务，包括交互式邮件卡片/详情卡片/分类选择卡片的设计与构建、卡片按钮事件处理与刷新链路，以及大模型分类/重要性/摘要的 Prompt 设计与参数调优、效果评估。"],
    ],
    widths=[55, 150, 300], size=10,
)
para("两名成员共同完成需求分析、系统联调、缺陷修复与文档撰写工作。", indent=True)

# ====================================================================
# 第2章 系统需求分析
# ====================================================================
page_break()
h1("第2章 系统需求分析")

h2("2.1 引言")
para("本章从用户角度出发，系统地分析 EmailClaw 智能邮件管家的功能需求与非功能需求，并对系统的目标用户、典型用例、业务流程进行建模，最后给出关键技术选型。需求分析是系统设计与实现的依据，本章在原始《EmailClaw 需求规格说明书》的基础上，结合实际实现情况进行了梳理与修订。需要特别说明的是，在实现过程中，项目对若干早期设计做了务实调整：交互与通知渠道由最初规划的 Telegram 调整为更易于企业内网联调、卡片交互能力更强的飞书（Lark）机器人；系统由最初设定的单用户演示形态演进为支持注册登录、凭据加密与数据隔离的多用户系统。下文均以实际实现为准进行描述。")

h2("2.2 功能需求")
para("EmailClaw 的目标用户主要包括三类：每天收发大量邮件的职场人士、需要区分学业通知与社交邮件的学生、以及关注邮件流量与分类统计的管理人员。围绕这些用户的核心诉求，系统定义了 FR-1 至 FR-7 共 7 项功能需求，下面分别说明。")

h3("FR-1　邮件实时接收与存储（多用户）")
para("系统能够为每个绑定邮箱的用户建立独立的 IMAP 长连接，实时监听收件箱（INBOX）的新邮件，抓取后通过 mailparser 解析 MIME 格式，提取发件人、收件人、主题、纯文本正文、HTML 正文与附件信息，并按用户隔离持久化到 PostgreSQL。连接断开时自动重连，并以自定义 IMAP 关键字与数据库标记双重去重，避免重复处理同一封邮件。", indent=True)

h3("FR-2　智能分类")
para("系统对每封新邮件进行自动分类，类别固定为 work（工作）、personal（个人）、shopping（购物）、marketing（营销）、spam（垃圾）、other（其他）六类。分类由两条路径完成：用户自定义规则引擎优先匹配，命中则直接给出分类；未命中则交由智能分析 Agent 处理。Agent 既可使用内置的本地关键词规则，也可在开启增强后调用 OpenClaw + DeepSeek 大模型，分类结果附带置信度与分类理由，便于用户理解判断依据。", indent=True)

h3("FR-3　重要性评分与主动推送")
para("系统为每封邮件计算 0～10 的重要性评分。用户可在偏好中设置重要性阈值（默认 7）与是否推送全部邮件（默认是）。当邮件重要性达到阈值时，推送的飞书卡片会以红色卡头并加“🔥”前缀高亮，提示用户优先处理；推送目标优先使用用户绑定的飞书 openId，未绑定时回退到机器人默认接收者。", indent=True)

h3("FR-4　邮件摘要生成")
para("系统自动为邮件生成简明摘要，帮助用户在不展开正文的情况下快速了解邮件主旨。本地规则模式下基于正文截取生成摘要；大模型增强模式下由 DeepSeek 生成更凝练的一句话摘要。摘要随分类结果一并展示在卡片上，并持久化到数据库。", indent=True)

h3("FR-5　用户交互与反馈（含规则自定义）")
para("用户通过飞书卡片按钮与系统交互，支持标为已读、标为重点、归档、删除、分类正确、分类错误（纠错）、查看详情、重新分析共 8 种操作。其中标记类操作会真实作用于 IMAP 邮箱；分类纠错会记录用户反馈并更新分类。此外，用户可通过 REST API 对自定义规则进行增删改查，规则以“条件—动作”形式描述，命中后可指定分类、重要性、摘要，并附带真实邮箱副作用动作。", indent=True)

h3("FR-6　邮件搜索与过滤")
para("用户可按分类、发件人、日期范围、重要性区间、读取/归档/删除状态以及关键词（在主题、发件人、正文中模糊匹配）等多条件组合检索历史邮件，结果分页返回。所有查询强制带用户标识，确保用户只能检索到属于自己的邮件。", indent=True)

h3("FR-7　每日摘要与统计")
para("规划中系统每日定时汇总当日邮件总数、未读数、分类分布、重要邮件列表等统计信息并推送给用户。该需求优先级为中，当前版本尚未实现，已列入后续工作。", indent=True)

para("各功能需求的优先级与实现状态汇总如下表所示。", indent=True)
add_table(
    ["需求编号", "需求名称", "优先级", "实现状态"],
    [
        ["FR-1", "邮件实时接收与存储（多用户）", "高", "已实现"],
        ["FR-2", "智能分类（规则 + 本地 Agent + 可选大模型）", "高", "已实现"],
        ["FR-3", "重要性评分与主动推送", "高", "已实现"],
        ["FR-4", "邮件摘要生成", "高", "已实现"],
        ["FR-5", "用户交互与反馈、规则自定义", "高", "已实现"],
        ["FR-6", "邮件搜索与过滤", "中", "已实现"],
        ["FR-7", "每日摘要与统计", "中", "未实现（后续工作）"],
    ],
    widths=[55, 280, 50, 110], size=9.5,
)
caption("表 2-1　功能需求追踪表")

h3("用例建模")
para("系统的主要参与者包括：用户（与系统交互的人）、邮箱服务器（IMAP 服务提供方）、飞书平台（通知与交互渠道）、数据库（数据存储）以及可选的 OpenClaw/大模型（分析增强）。围绕参与者抽象出的核心用例及其关系如下图所示。", indent=True)
code_block(r"""
                        ┌─────────────────────┐
   邮箱服务器(IMAP) ───▶│  UC1 邮箱连接/鉴权    │
                        │  UC2 实时接收新邮件   │
                        │  UC3 邮件解析与入库   │
                        └──────────┬──────────┘
                                   ▼
              ┌───────────────────────────────────────┐
              │  UC4 规则匹配      UC5 智能分类(分类/   │
              │  (命中跳过Agent)   重要性/摘要)         │◀── OpenClaw/DeepSeek
              └──────────┬───────────────┬─────────────┘
                         ▼               ▼
                 ┌───────────────┐  ┌──────────────────┐
   用户 ────────▶│ UC6 重要性判断 │  │ UC7 飞书卡片推送 │──▶ 飞书平台
     ▲           │ 与推送决策     │  └────────┬─────────┘
     │           └───────────────┘           ▼
     │                              ┌──────────────────────┐
     └──────────────────────────────│ UC8 卡片交互(已读/重点│
        UC9 邮件搜索/UC10 规则管理   │ /归档/删除/纠错/详情/ │──▶ IMAP 真实操作
                                     │ 重新分析)            │──▶ 数据库同步
                                     └──────────────────────┘
""")
caption("图 2-1　EmailClaw 系统用例图")

para("以最核心的 UC5（智能分类）为例：前置条件为邮件已成功接收并入库；主流程为系统提取邮件主题、发件人与正文 → 先经规则引擎匹配 → 未命中则提交 Agent 分析 → Agent（本地规则或大模型）综合判断返回 {分类, 置信度, 理由, 重要性, 摘要} → 系统持久化分类结果并更新邮件字段；备选流程为大模型分析失败或超时时，系统记录日志并自动回退本地规则 Agent，保证分类不中断；后置条件为邮件的分类、重要性、摘要字段已更新且分类记录已落库。", indent=True)

h2("2.3 非功能需求")
para("除功能性需求外，系统还需满足性能、安全、可靠、可维护与用户体验等非功能性需求。", indent=True)

h3("性能需求")
add_table(
    ["指标", "目标值", "说明"],
    [
        ["邮件接收延迟", "秒级", "IMAP 推送结合 30s 兜底轮询，新邮件到达后秒级被抓取"],
        ["分析延迟", "本地 < 1s / 大模型 < 30s", "本地规则毫秒级；大模型受 DeepSeek 网络与推理耗时影响"],
        ["卡片交互响应", "即时（先 toast）", "先返回 toast 规避飞书 3 秒回调超时，再异步刷新卡片"],
        ["邮件搜索", "< 1s（千封内）", "数据库索引 + 分页查询"],
        ["并发处理", "多用户并发", "每用户独立 IMAP 连接与子进程隔离"],
    ],
    widths=[110, 150, 250], size=9.5,
)
caption("表 2-2　性能需求")

h3("安全需求")
add_table(
    ["需求", "实现方式"],
    [
        ["身份认证", "邮箱 + 密码注册登录，密码经 bcrypt（12 轮）哈希，签发 JWT（7 天有效期）"],
        ["敏感信息加密存储", "IMAP 授权码使用 AES-256-GCM 加密入库，密钥经 ENCRYPTION_KEY 派生"],
        ["访问控制", "所有用户 API 强制 JWT；查询与写操作强制按 userId 过滤与归属校验"],
        ["内部服务鉴权", "飞书 bot → 后端的 webhook 调用必须携带共享密钥 X-Bot-Secret"],
        ["操作审计", "AgentLog 表记录每次分析的输入、输出、模型与耗时"],
    ],
    widths=[120, 390], size=9.5,
)
caption("表 2-3　安全需求")

h3("可靠性、可维护性与用户体验需求")
bullet("可靠性：IMAP 断线自动重连；大模型分析失败自动回退本地规则；邮件处理主流程对推送失败做容错，不阻塞后续；以数据库 notifiedAt 实现跨重启的持久化去重。")
bullet("可维护性：代码按职责清晰分层（路由 / 服务 / 集成 / 工具），关键日志带 [user=...] 前缀；TypeScript 静态类型检查；数据库 schema 配套数据字典。")
bullet("用户体验：飞书卡片格式美观，重要邮件高亮，删除操作二次确认，操作后即时反馈并刷新卡片状态，失败有明确提示。")

h2("2.4 技术选型")
para("综合考虑开发效率、生态成熟度与项目实际需求，系统的关键技术选型如下表所示。", indent=True)
add_table(
    ["类别", "技术 / 框架", "选型说明"],
    [
        ["语言/运行时", "Node.js 20+ / TypeScript 5", "异步 I/O 适合邮件与网络密集型场景；静态类型提升可维护性"],
        ["Web 框架", "Express 5", "轻量成熟，便于快速构建 REST API"],
        ["ORM/数据库", "Prisma 7 + PostgreSQL", "类型安全的数据访问，JSON 字段支持灵活存储偏好与规则"],
        ["邮件协议", "imap + mailparser", "IMAP 长连接监听 + MIME 解析"],
        ["鉴权/加密", "bcryptjs / jsonwebtoken / AES-256-GCM", "密码哈希、令牌签发、IMAP 凭据加密"],
        ["即时通讯交互", "飞书开放平台 Node SDK", "交互式卡片能力强，WebSocket 长连接接收事件，企业内网联调便捷"],
        ["AI 分析（可选）", "OpenClaw + DeepSeek", "OpenClaw 作为大模型编排器按需拉起；DeepSeek 成本低、中文效果好"],
    ],
    widths=[90, 200, 220], size=9.5,
)
caption("表 2-4　关键技术选型")

h2("2.5 其他要求")
bullet("运行环境：Node.js 18+（建议 20+）、可访问的 PostgreSQL、支持 IMAP 的邮箱及其授权码；如需飞书交互需配置企业自建应用；如需大模型增强需安装 OpenClaw CLI 并配置 DeepSeek API Key。")
bullet("部署形态：后端服务（默认端口 3000）与飞书机器人服务（默认端口 3001）双进程协作，可分别独立启动与调试。")
bullet("约束：邮件分类固定为 6 类；删除操作为真实不可恢复删除，需谨慎；状态同步为单向（EmailClaw → 邮箱服务器）。")

# ====================================================================
# 第3章 系统设计
# ====================================================================
page_break()
h1("第3章 系统设计")

h2("3.1 引言")
para("本章在需求分析的基础上，自顶向下地阐述 EmailClaw 的系统设计：首先给出功能模块划分，其次说明整体架构与端到端数据流，再对飞书交互卡片进行 UI 设计，随后对各核心模块进行详细设计（数据结构、接口与关键算法），最后给出数据库设计与数据字典。", indent=True)

h2("3.2 系统功能设计")
para("系统按职责划分为五大功能模块，模块结构如下图所示。", indent=True)
code_block(r"""
                         EmailClaw 智能邮件管家
                                 │
   ┌───────────────┬──────────────┼───────────────┬───────────────┐
   ▼               ▼              ▼               ▼               ▼
①用户与鉴权    ②邮件接入与解析  ③智能分析       ④交互与推送      ⑤检索与管理
 ·注册/登录     ·多用户IMAP连接  ·规则引擎        ·飞书卡片推送     ·邮件多条件搜索
 ·JWT鉴权       ·断线重连        ·本地规则Agent   ·8种按钮交互      ·邮件详情
 ·IMAP凭据加密  ·MIME解析        ·OpenClaw+大模型 ·先toast再刷新    ·规则CRUD
 ·三层防御鉴权  ·去重(标记+DB)   ·失败自动回退    ·真实IMAP同步     ·用户偏好设置
""")
caption("图 3-1　系统功能模块图")

h2("3.3 系统架构与数据流设计")
para("系统采用双服务协作的架构：主后端服务（端口 3000）负责邮件抓取解析、智能分析、数据持久化、REST API 与真实邮箱操作；飞书机器人服务（端口 3001）负责构建并发送交互式卡片、通过 WebSocket 长连接接收卡片按钮事件，并将事件转发回后端。两者通过 HTTP 互相调用，并以共享密钥相互鉴权。整体架构与一封新邮件的完整数据流如下图所示。", indent=True)
code_block(r"""
 ┌────────────┐  IMAP长连接   ┌──────────────────────────────────────────┐
 │ 用户邮箱    │◀────────────▶│            主后端服务 (:3000)              │
 │(QQ/Gmail..) │  抓取/操作     │  UserMailbox ── ImapManager                │
 └────────────┘               │       │ onIncomingEmail                     │
                              │       ▼                                     │
                              │  processIncomingEmail 流水线:               │
                              │   1.upsertEmail 入库                        │
                              │   2.notifiedAt 去重                         │
   ┌────────────┐  Prisma     │   3.ruleEngine 规则匹配(命中跳过Agent)      │
   │ PostgreSQL │◀───────────▶│   4.agentService 分析(本地/OpenClaw)        │
   └────────────┘             │   5.feishuService 推送决策(阈值/高亮)       │
                              │   6.markNotified 标记已处理                 │
                              └───────────┬───────────────────▲────────────┘
                                  POST /api/notify-email   POST /api/feishu/webhook
                                          ▼                   │ (X-Bot-Secret)
                              ┌──────────────────────────────────────────┐
                              │          飞书机器人服务 (:3001)            │
                              │  buildEmailCard → feishuClient.sendCard    │
                              │  WSClient 监听 card.action.trigger         │
                              │  先 toast → forwardToBackend → updateCard  │
                              └───────────┬──────────────────────────────┘
                                          ▼  飞书 OpenAPI
                                   ┌─────────────┐
                                   │  飞书用户    │  点击按钮 ◀───────┐
                                   └─────────────┘                   │
                                          └─── 已读/重点/归档/删除…──┘
""")
caption("图 3-2　系统架构与端到端数据流图")

para("推送链路：后端分析完成后，按用户偏好（阈值、是否全量推送）决定是否推送，调用飞书 bot 的 /api/notify-email，bot 据此构建卡片并通过飞书 OpenAPI 发送。回调链路：用户点击卡片按钮触发 card.action.trigger 事件，bot 先同步返回轻量 toast 让飞书停止 3 秒倒计时，再在后台携带共享密钥转发到后端 /api/feishu/webhook；后端经“共享密钥 → openId 解析 userId → emailId 归属”三重校验后执行业务，并把最新邮件状态回传，bot 主动 patch 原卡片完成刷新。", indent=True)

h2("3.4 系统 UI 设计")
para("系统的用户界面以飞书交互式卡片为载体，包含三类卡片：邮件通知卡片、分类选择卡片、邮件详情卡片。邮件通知卡片是核心界面，其布局示意如下。", indent=True)
code_block(r"""
┌──────────────────────────────────────────────┐
│ 🔥 邮件主题            （重要邮件红色卡头+🔥）  │   ← header(按分类着色)
├──────────────────────────────────────────────┤
│ 发件人：xxx           收件人：yyy              │
│ 时间：2026-06-06 10:00                          │
│ ──────────────────────────────────────────    │
│ 📂 分类：工作   置信度 90%   ⭐重要性 ★★★★★★★★ (8/10) │
│ ──────────────────────────────────────────    │
│ 📝 摘要：……                                    │
│ 💡 分类理由：……                                │
│ 状态：📬未读 | 📂收件箱                          │
│ ──────────────────────────────────────────    │
│ [📖标为已读] [⭐标为重点] [📦归档]             │
│ [✅分类正确] [❌分类错误] [🔍查看详情]         │
│ [🗑删除]    [🔄重新分析]                        │
└──────────────────────────────────────────────┘
""")
caption("图 3-3　飞书邮件通知卡片布局")

para("卡片以分类映射卡头颜色，便于快速区分邮件类型；当邮件被判定为重要时统一覆盖为红色并加“🔥”前缀。分类与颜色映射、按钮行为分别如下两表。", indent=True)
add_table(
    ["分类", "work", "personal", "shopping", "marketing", "spam", "other"],
    [["中文", "工作", "个人", "购物", "营销", "垃圾", "其他"],
     ["卡头色", "蓝", "绿", "黄", "紫", "红", "灰"]],
    size=9.5,
)
caption("表 3-1　分类与卡片卡头颜色映射")

add_table(
    ["按钮", "动作标识", "行为"],
    [
        ["标为已读", "mark_read", "IMAP 加 \\Seen 标志，数据库 isRead=true，异步刷新卡片"],
        ["标为重点", "mark_important", "IMAP 加 \\Flagged 标志，数据库 importance=10"],
        ["归档", "archive", "IMAP MOVE 到归档目录，数据库 isArchived=true"],
        ["删除", "delete", "IMAP \\Deleted + expunge（永久删除），数据库 isDeleted=true，按钮区消失"],
        ["分类正确", "feedback_correct", "记录 Classification.feedback='correct'"],
        ["分类错误", "feedback_wrong", "弹出分类选择卡片，选择后写入正确分类并刷新"],
        ["查看详情", "view_detail", "新发一张详情卡片，保留原卡片按钮"],
        ["重新分析", "reanalyze", "仅重跑 Agent（不重走规则），更新并刷新当前卡片"],
    ],
    widths=[60, 95, 355], size=9.5,
)
caption("表 3-2　卡片按钮行为设计")

h2("3.5 系统详细设计")

h3("3.5.1 多用户 IMAP 连接管理")
para("为支持多用户，系统将 IMAP 连接抽象为两层：UserMailbox 封装单个用户的 IMAP 长连接与邮件操作，ImapManager 作为多用户连接注册表统一调度。每个绑定邮箱的用户对应一个 UserMailbox 实例，连接彼此隔离，所有日志带 [user=...] 前缀以便追踪。", indent=True)
bullet("连接与重连：建立 TLS 长连接并开启 keepalive；监听 error/end 事件，断开后延迟 5 秒自动重连。")
bullet("新邮件感知：连接就绪后打开 INBOX，先全量扫描一次未读邮件，再监听 mail 事件；考虑到 QQ 等邮箱 IMAP IDLE 推送不可靠，额外加入 30 秒兜底轮询。")
bullet("解析与回调：对未读邮件按 UID 抓取，用 mailparser 解析为统一的 SimpleEmail 结构，回调交给上层处理流水线。")
bullet("去重设计：服务器侧写入自定义关键字 CLAWED 标记已处理；同时维护内存级 claimedUids 集合与 scanning 重入锁，挡住 CLAWED 写入前的并发重复扫描；考虑到部分邮箱不持久化自定义关键字，最终以数据库 notifiedAt 字段作为跨重启去重的权威依据。")
bullet("邮箱配置推断：resolveImapConfig 根据邮箱域名自动推断常见服务商（QQ/163/Gmail/Outlook 等）的 IMAP 主机与端口；归档目录不存在时自动创建。")
para("UserMailbox 对外提供 markRead、markFlagged、archive、deleteMail 等操作，均以 Promise 封装 IMAP 异步回调；ImapManager 提供 startForAllBoundUsers、startForUser、stopForUser、getMailboxForEmail 等接口，其中 getMailboxForEmail 在按 emailId 路由到对应连接时会再次做归属校验。", indent=True)

h3("3.5.2 邮件处理流水线")
para("新邮件统一进入 EmailService.processIncomingEmail 流水线，串联入库、去重、规则、分析、推送与标记六个步骤。其核心逻辑可用如下伪代码表示。", indent=True)
code_block(r"""
processIncomingEmail(userId, email):
  saved = upsertEmail(userId, email)          # 1. 入库(按 userId+uid 唯一)
  if saved.notifiedAt: return                  # 2. 已处理过 → 跳过(跨重启去重)
  ruleHit = ruleEngine.evaluate(userId, email) # 3. 先跑规则
  if ruleHit:
      analysis = ruleHit.result                #    命中规则 → 跳过 Agent
      applyRuleSideEffects(ruleHit.sideEffects)#    执行 mark_read/archive/delete
  else:
      analysis = agentService.analyzeEmail()   # 4. 未命中 → Agent 分析
  threshold = user.preferences.importanceThreshold ?? 7
  isImportant = analysis.importance >= threshold
  if pushAllEmails or isImportant:             # 5. 按偏好决定是否推送
      feishuService.pushEmailCard(..., isImportant)
  markNotified(saved.id)                        # 6. 标记已处理
""")
caption("图 3-4　邮件处理流水线伪代码")

h3("3.5.3 规则引擎")
para("规则引擎在 Agent 之前运行，命中即跳过 Agent，从而让用户对确定性场景拥有更高优先级、更可控的处理方式。规则存储于 Rule 表，conditions 与 actions 均为 JSON，结构约定如下。", indent=True)
code_block(r"""
conditions: [                         // 多个条件之间为 AND 关系
  { field: from|to|subject|body,
    operator: contains|equals|startsWith|endsWith|regex,
    value: "...", caseSensitive?: false }
]
actions: {
  category: work|personal|shopping|marketing|spam|other,  // 必填
  importance?: 0-10,                  // 默认 5
  summary?: "自定义摘要",             // 默认走自动摘要
  sideEffects?: [mark_read|archive|delete]   // 命中后真实 IMAP 副作用
}
""")
para("引擎按 priority 降序遍历用户启用的规则，对每条规则要求其全部条件命中（every）；首个命中的规则合成 EmailAgentResult 返回，分类置信度记为 1、模型标记为 rule-engine-v1，并携带需要执行的副作用动作。引擎另提供 validateDefinition 用于创建/更新规则时的字段校验（field/operator 合法性、regex 可编译性、importance 取值范围等）。", indent=True)

h3("3.5.4 智能分析 Agent")
para("AgentService 是分析层的统一入口，其 analyzeEmailDraft 优先尝试 OpenClaw（若启用），失败则回退到本地规则 Agent；analyzeEmail 在此基础上额外持久化分析结果与 AgentLog。本地规则 Agent 由 ClassificationSkill 实现，包含三项能力：", indent=True)
bullet("分类：将主题、发件人、收件人、正文与附件名归一化为小写文本，按 spam→work→shopping→marketing→personal 的优先顺序进行关键词匹配，命中即定类，否则归为 other；命中类别置信度 0.72，other 为 0.5。")
bullet("重要性评分：基础分 3，命中“紧急/截止/会议/deadline/urgent”等关键词 +4，含附件 +1，上限 10。")
bullet("摘要：截取正文前若干字符并拼接发件人与主题，生成一句话摘要。")
para("这种本地规则方式不依赖任何外部服务即可跑通完整流程，作为大模型不可用时的兜底，保证系统始终可用。", indent=True)

h3("3.5.5 OpenClaw + DeepSeek 大模型接入")
para("当配置 OPENCLAW_ENABLED=true 时，OpenClawClient 会为每封邮件临时拉起一个 OpenClaw 子进程，以本机内嵌（--local）模式调用 DeepSeek 完成分析，命令形如 openclaw agent --agent email-claw --session-key <唯一> --message <提示词> --json --local。由于大模型与 OpenClaw 返回结构的不稳定性，客户端做了三处关键适配：", indent=True)
bullet("拆信封解析：OpenClaw 的 --json 返回 {payloads, meta} 信封，真正的分析 JSON 套在 payloads[0].text 字符串里，需先拆出再解析，否则会静默拿到默认值（每封都退化为 other / 重要性 3）。")
bullet("独立 session-key：后端会并发拉起多个 --local 子进程，共用默认会话会触发会话文件锁抢占错误（TakeoverError），故为每封邮件生成唯一 session-key 实现隔离。")
bullet("容错与重试：DeepSeek 偶发会把 JSON 包进 ```json``` 围栏或输出非法 JSON，客户端会剥除围栏、用正则兜底提取，并在首次失败后换 session-key 重试一次，两次都失败才抛出由上层回退本地规则。")
para("分析结果经 normalizeResult 归一化：分类落到 6 类合法值之外时归为 other，置信度裁剪到 [0,1]，重要性裁剪到 [0,10]，缺省字段填充默认值，确保下游拿到结构稳定的结果。", indent=True)

h3("3.5.6 飞书交互服务")
para("FeishuService 负责推送与回调两件事。pushEmailCard 将分析结果 POST 到飞书 bot 的 /api/notify-email，并对失败做 try/catch 吞错以避免阻塞邮件主流程。handleCallback 是回调分发中枢：先做 emailId 归属校验，再按 action 分发到 8 个处理方法，对标记类操作先调用 EmailService 执行真实 IMAP 操作、再更新数据库、最后回传最新邮件状态。飞书 bot 端的 cardHandler 采用“先 toast 再异步 updateCard”模式：同步返回“处理中…”toast 让飞书停止倒计时，后台 await 后端结果后再 patch 原卡片；查看详情则新发一张详情卡片。这一设计根治了早期“为规避 3 秒超时改为立即返回空响应、导致卡片永不刷新”的缺陷。", indent=True)

h3("3.5.7 防御性鉴权三层防线")
para("针对多账号邮件场景的安全需求，系统设计了 HTTP、业务、数据三层防御性鉴权，层层校验，确保用户无法越权访问或操作他人邮件。", indent=True)
add_table(
    ["层级", "位置", "校验内容"],
    [
        ["HTTP 层", "requireAuth / requireBotSecret 中间件", "所有用户 API 强制 JWT；飞书 webhook 必须携带正确的 X-Bot-Secret"],
        ["业务层", "feishuRoutes", "先将飞书 openId 解析为 userId，未绑定飞书的请求直接 403"],
        ["数据层", "assertEmailOwnership / updateMany where userId", "每次写操作前校验 emailId 归属；规则增删改强制带 userId 防越权"],
    ],
    widths=[55, 220, 235], size=9.5,
)
caption("表 3-3　防御性鉴权三层防线")
para("配套的安全工具包括：crypto 模块用 AES-256-GCM 加解密 IMAP 授权码（密钥由 ENCRYPTION_KEY 经 SHA-256 派生，密文为 iv‖密文‖认证标签 的 base64）；auth 模块用 bcrypt（12 轮）哈希密码、用 jsonwebtoken 签发与校验 7 天有效期的 JWT。", indent=True)

h3("3.5.8 REST API 设计")
para("后端对外提供认证、用户自管理、邮件检索、规则管理与飞书回调五组 API，主要接口如下表（标注 [JWT] 的需登录令牌）。", indent=True)
add_table(
    ["方法", "路径", "说明"],
    [
        ["POST", "/api/auth/register", "注册（密码≥8 位），返回 token 与用户信息"],
        ["POST", "/api/auth/login", "登录，返回 token 与用户信息"],
        ["GET", "/api/users/me [JWT]", "获取当前用户信息、偏好与邮箱绑定状态"],
        ["PATCH", "/api/users/me/preferences [JWT]", "设置重要性阈值(1-10)与是否推送全部邮件"],
        ["PATCH", "/api/users/me/feishu [JWT]", "绑定飞书 openId"],
        ["POST", "/api/users/me/mailbox [JWT]", "绑定 IMAP（加密入库并即时启动连接）"],
        ["DELETE", "/api/users/me/mailbox [JWT]", "解绑邮箱并断开 IMAP"],
        ["GET", "/api/emails [JWT]", "多条件分页搜索邮件"],
        ["GET", "/api/emails/:id [JWT]", "单封详情（自动校验归属）"],
        ["GET/POST/PATCH/DELETE", "/api/rules [JWT]", "规则的增删改查（按优先级降序）"],
        ["POST", "/api/feishu/webhook", "飞书卡片按钮事件回调（需 X-Bot-Secret）"],
    ],
    widths=[95, 200, 215], size=9,
)
caption("表 3-4　REST API 一览")

h2("3.6 数据库设计")
para("系统使用 PostgreSQL，通过 Prisma 进行建模与访问，共设计 6 张数据表：User（用户）、Email（邮件）、Classification（分类结果）、Contact（联系人）、Rule（规则）、AgentLog（Agent 执行日志）。各表关系如下图所示。", indent=True)
code_block(r"""
        User (1) ───< (N) Email (1) ───< (1) Classification
          │  ├──< (N) Rule              (每封邮件一条分类结果)
          │  ├──< (N) Contact
          │  └──< (N) AgentLog
   说明: Email 以 (userId, uid) 唯一; 索引(userId, receivedAt)、(category)
""")
caption("图 3-5　数据库 ER 关系图")

para("核心表 User、Email、Classification 的数据字典如下（其余表结构类似，从略）。", indent=True)
add_table(
    ["字段", "类型", "说明"],
    [
        ["id", "UUID(PK)", "用户 ID"],
        ["email", "String 唯一", "登录邮箱"],
        ["password", "String", "bcrypt 哈希后的密码"],
        ["imapHost/imapUser", "String?", "IMAP 服务器与登录名"],
        ["imapPassword", "String?", "AES-256-GCM 加密后的 IMAP 授权码"],
        ["feishuUserId", "String?", "绑定的飞书 openId"],
        ["preferences", "JSON", "偏好：importanceThreshold、pushAllEmails 等"],
        ["createdAt/updatedAt", "DateTime", "创建/更新时间"],
    ],
    widths=[120, 110, 280], size=9.5,
)
caption("表 3-5　User 表数据字典")
add_table(
    ["字段", "类型", "说明"],
    [
        ["id", "UUID(PK)", "邮件 ID"],
        ["userId", "FK", "所属用户"],
        ["messageId / uid", "String / Int", "IMAP 消息标识与 UID"],
        ["from/to/cc/subject", "String", "发件人/收件人/抄送/主题"],
        ["body / html", "Text", "纯文本正文 / HTML 正文"],
        ["category", "String?", "分类：work/personal/shopping/marketing/spam/other"],
        ["importance", "Int?", "重要性 0-10"],
        ["summary", "Text?", "邮件摘要"],
        ["isRead/isArchived/isDeleted", "Boolean", "已读/归档/删除状态"],
        ["notifiedAt", "DateTime?", "已完成处理时间，用于跨重启持久化去重"],
        ["receivedAt", "DateTime", "邮件接收时间"],
    ],
    widths=[150, 100, 260], size=9.5,
)
caption("表 3-6　Email 表数据字典（唯一约束 userId+uid）")
add_table(
    ["字段", "类型", "说明"],
    [
        ["id", "UUID(PK)", "分类记录 ID"],
        ["emailId", "FK 唯一", "对应邮件"],
        ["userId", "FK", "所属用户"],
        ["category / confidence", "String / Float", "分类结果与置信度"],
        ["reasoning", "Text?", "分类理由（含重要性与摘要拼接）"],
        ["toolsUsed / executionSteps", "JSON?", "使用的工具与执行步骤"],
        ["model", "String", "模型标识：local-rule-agent-v0 / openclaw-agent / rule-engine-v1"],
        ["feedback", "String?", "用户反馈：correct / incorrect"],
    ],
    widths=[150, 110, 250], size=9.5,
)
caption("表 3-7　Classification 表数据字典")
para("此外，AgentLog 表记录每次分析的类型、状态、输入输出、模型与耗时，用于操作审计与问题排查；Contact 表为后续基于发件人历史的重要性评估预留；Rule 表存储用户自定义规则。", indent=True)

# ====================================================================
# 第4章 系统实现
# ====================================================================
page_break()
h1("第4章 系统实现")

h2("4.1 引言")
para("本章介绍系统的实现情况，包括代码组织、典型运行界面与调用示例，并给出测试计划与测试报告。系统在 Linux（含 WSL2）环境下开发，后端使用 ts-node-dev 热重载启动，TypeScript 编译验证通过，飞书 bot 关键脚本语法检查通过。源码按职责清晰分层，主要目录结构如下。", indent=True)
code_block(r"""
src/
├── server.ts                         后端入口(:3000)，挂载路由 + bootstrap
├── backend/
│   ├── api/routes/                   认证/用户/邮件/规则/飞书 回调路由
│   ├── middleware/authMiddleware.ts  requireAuth / requireBotSecret
│   ├── services/
│   │   ├── emailService.ts           邮件处理流水线 facade
│   │   ├── userMailbox.ts            单用户 IMAP 长连接封装
│   │   ├── imapManager.ts            多用户 IMAP 连接注册表
│   │   ├── ruleEngine.ts             规则匹配引擎
│   │   ├── agentService.ts           本地 Agent + OpenClaw 调度
│   │   ├── openClawClient.ts         OpenClaw/DeepSeek 子进程适配
│   │   ├── databaseService.ts        Prisma 数据访问
│   │   └── bootstrap.ts              启动自检 / 引导用户
│   ├── integrations/feishu/feishuService.ts   推送 + 回调分发
│   └── utils/{auth.ts, crypto.ts}    JWT/bcrypt、AES-256-GCM
├── agents/                           本地规则 Agent、Skill 与工具
email_claw_bot/src/
├── server.js                         飞书 bot HTTP + WebSocket(:3001)
├── cards/emailCard.js                邮件/详情/分类选择 卡片构建器
├── handlers/cardHandler.js           按钮事件转发 + updateCard 刷新
└── services/feishuClient.js          飞书 OpenAPI 封装
prisma/schema.prisma                  6 张表数据模型 + 迁移
""")
caption("图 4-1　项目源码目录结构")

h2("4.2 系统典型界面")
para("系统的典型“界面”由两部分组成：面向终端用户的飞书交互式卡片，以及面向开发者/前端的 REST API。", indent=True)
h3("飞书交互卡片")
para("用户在飞书中收到的邮件通知卡片如 3.4 节图 3-3 所示：卡头按分类着色、重要邮件红色高亮，卡片体依次展示发件人/收件人、时间、分类与置信度、星级重要性、摘要、分类理由与当前状态，底部为三行共 8 个操作按钮。点击“查看详情”会新发一张包含正文预览的详情卡片；点击“分类错误”会弹出 6 个分类选项的选择卡片；其余操作执行后会原地刷新卡片状态，并附一条带邮件主题的简短确认消息。删除按钮带二次确认弹窗，提示该操作将真实且不可恢复地删除邮件。", indent=True)
h3("REST API 调用示例")
para("以注册、登录、绑定邮箱与搜索邮件为例，典型调用如下。", indent=True)
code_block(r"""
# 注册
curl -X POST :3000/api/auth/register -H 'Content-Type: application/json' \
  -d '{"email":"u@example.com","password":"pass1234"}'
# → { token, user }

# 绑定 IMAP 邮箱(需登录令牌)
curl -X POST :3000/api/users/me/mailbox -H "Authorization: Bearer $TOKEN" \
  -H 'Content-Type: application/json' \
  -d '{"imapHost":"imap.qq.com","imapUser":"x@qq.com","imapPassword":"授权码"}'
# → 加密入库并后台启动该用户 IMAP 连接

# 多条件搜索邮件
curl ":3000/api/emails?category=work&importanceMin=7&q=会议&page=1" \
  -H "Authorization: Bearer $TOKEN"
# → { items, total, page, pageSize, totalPages }

# 健康检查
curl :3000/ping   # {"message":"🏓 EmailClaw 已启动"}
curl :3001/ping   # {"status":"ok","service":"EmailClaw Bot"}
""")
caption("图 4-2　REST API 调用示例")

h2("4.3 测试计划")
para("测试由项目两名成员承担，采用功能联调测试与安全（越权）测试相结合的方式。测试环境为：Linux/WSL2 + Node.js 20 + 本地 PostgreSQL；邮箱使用 QQ 邮箱（IMAP 授权码）；飞书使用企业自建应用；大模型增强使用 OpenClaw CLI + DeepSeek（deepseek-chat）。主要测试用例如下表所示。缺陷管理上，联调中发现的问题以提交记录跟踪，修复后回归验证。", indent=True)
add_table(
    ["编号", "测试项", "测试要点", "预期结果"],
    [
        ["TC-01", "邮件抓取入库", "向绑定邮箱发测试邮件", "秒级抓取、解析并入库，不重复"],
        ["TC-02", "本地规则分类", "关闭大模型，发不同类型邮件", "正确分入 work/marketing/other 等"],
        ["TC-03", "大模型增强分析", "开启 OpenClaw+DeepSeek", "分类非清一色 other、重要性非清一色 3"],
        ["TC-04", "失败回退", "断开/禁用大模型", "自动回退本地规则，仍正常收邮件"],
        ["TC-05", "规则命中", "配置规则并发匹配邮件", "命中规则、跳过 Agent、执行副作用"],
        ["TC-06", "重要性推送", "发高/低重要性邮件", "达阈值红色高亮推送，未达按偏好处理"],
        ["TC-07", "卡片交互", "依次点击 8 个按钮", "真实同步 IMAP+DB，卡片正确刷新"],
        ["TC-08", "邮件搜索", "多条件组合查询", "结果正确、分页正常、< 1s"],
        ["TC-09", "认证与鉴权", "注册/登录/JWT 校验", "无令牌或令牌无效被拦截"],
        ["TC-10", "越权防御", "用他人令牌读/操作邮件", "被三层防线拦截"],
    ],
    widths=[45, 95, 175, 195], size=9,
)
caption("表 4-1　主要测试用例")

h2("4.4 测试报告")
para("在完成上述实现后，项目对所有服务进行了真实启动与端到端联调，主要结果如下。", indent=True)
bullet("功能联调：引导用户自动建成并加密绑定邮箱；IMAP 鉴权成功并自动创建 Archive 归档目录；未读邮件被抓取、解析、入库并完成分析；登录获取 JWT 后，/api/users/me、/api/emails 搜索、/api/rules 增删改查全部通过；飞书卡片推送、8 类按钮交互与卡片刷新均正确，标记/归档/删除操作真实同步到邮箱客户端。")
bullet("大模型增强验证：按 OpenClaw + DeepSeek 接入文档完成配置后，分类结果不再退化为清一色 other、重要性不再清一色 3，日志无 TakeoverError 与解析失败；将 OPENCLAW_ENABLED 置为 false 后重启，系统正确回退本地规则且仍能正常收发。")
bullet("安全（越权）测试：攻击者用自身 JWT 读取他人 emailId 返回“无权操作此邮件”；列出他人邮件返回空集；未携带 X-Bot-Secret 调用 webhook 返回 401；未绑定飞书的 openId 调用 webhook 返回 403。三层防线全部命中。")
para("联调过程中发现并修复了两个关键缺陷，跟踪情况如下表。", indent=True)
add_table(
    ["缺陷", "现象", "根因", "修复"],
    [
        ["卡片不刷新", "查看详情无响应、各操作后卡片状态不变",
         "为规避飞书 3 秒超时，webhook 改为立即返回空响应，bot 拿不到结果",
         "后端恢复同步语义返回完整结果；bot 改为先 toast 再异步 updateCard"],
        ["重复发卡", "同一封邮件被多次推送飞书卡片",
         "IMAP 轮询并发 + CLAWED 标记写入存在时间窗，跨重启又会重扫",
         "新增内存级 claimedUids 锁与 scanning 重入锁，并以数据库 notifiedAt 做权威去重"],
    ],
    widths=[60, 130, 160, 160], size=9,
)
caption("表 4-2　缺陷跟踪与修复")
para("综合测试结果，系统已稳定打通“邮件到达 → IMAP 抓取 → MIME 解析 → 入库 → 规则/Agent 分析 → 飞书卡片展示 → 用户交互 → 真实邮箱与数据库同步 → 卡片刷新”的完整闭环，覆盖了 FR-1 至 FR-6 的核心功能。", indent=True)

# ====================================================================
# 第5章 总结
# ====================================================================
page_break()
h1("第5章 总结")

para("本项目设计并实现了一套面向多用户的智能邮件管家系统 EmailClaw，达成了课题既定目标。系统以 Node.js + TypeScript 为基础，构建了由后端服务与飞书机器人服务协作的双进程架构，完整实现了从邮件实时接收、解析入库，到规则与 AI 分析、重要性评分与摘要，再到飞书交互式卡片推送、用户反馈与真实邮箱操作的端到端闭环。", indent=True)

para("回顾需求规格说明书，FR-1（邮件实时接收与存储，并扩展为多用户）、FR-2（智能分类）、FR-3（重要性评分与主动推送）、FR-4（摘要生成）、FR-5（用户交互与反馈、规则自定义）、FR-6（邮件搜索与过滤）均已实现；仅 FR-7（每日摘要与统计）作为中优先级需求暂未完成。", indent=True)

para("本项目的主要技术亮点包括：", indent=True)
bullet("分层智能策略：规则引擎 + 本地关键词 Agent + 可选 OpenClaw/DeepSeek 大模型增强 + 失败自动回退，在成本、稳定性与效果之间取得务实平衡，系统在任何情况下都可用。")
bullet("健壮的多用户 IMAP 管理：连接隔离、断线重连、IDLE 与轮询双触发，并通过内存锁 + 数据库 notifiedAt 双重机制根治了并发与跨重启的重复处理问题。")
bullet("三层防御性鉴权：HTTP（JWT / 共享密钥）、业务（openId→userId）、数据（emailId 归属）层层校验，配合 bcrypt 与 AES-256-GCM，保障了多账号场景的数据隔离与安全。")
bullet("良好的交互体验：飞书卡片“先 toast 再异步刷新”模式既规避了 3 秒回调超时，又保证了操作后卡片状态的及时刷新与失败可见。")

para("项目同时也存在不足，后续可在以下方向继续完善：一是真正打通 Agent 反馈学习闭环，将用户的分类纠错回喂给分析模型以持续优化；二是实现 FR-7 的每日摘要与统计定时任务；三是补充面向终端用户的 Web 前端入口与自动化测试覆盖；四是进一步打磨大模型 Prompt 与生产级部署的稳定性。", indent=True)

para("在分工协作方面，两名成员各有侧重又紧密配合：后端服务与 OpenClaw 接入工作量大、技术难点集中在多用户并发、IMAP 协议细节与大模型子进程适配上，构成了系统的骨架与底座；前端飞书服务与模型参数调优工作直接决定了系统的可用体验与分析效果，对交互闭环的打通和分类质量的提升贡献显著。两人共同完成了需求分析、联调、缺陷修复与文档撰写。整体协作高效，分工合理，各成员均较好地完成了承担的任务，对项目的最终落地做出了实质性贡献。", indent=True)

para("通过本次综合项目实践，团队完整经历了从需求分析、系统设计、编码实现到测试联调的软件工程全流程，在网络协议、数据库、并发处理、第三方平台对接、AI Agent 编排与系统安全等方面都得到了切实锻炼，对工程化落地一个贴近真实形态的系统有了更深入的理解。", indent=True)

# ====================================================================
# 保存
# ====================================================================
doc.save(OUTPUT)
print("报告已生成：", OUTPUT)

