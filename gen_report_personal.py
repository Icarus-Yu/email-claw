# -*- coding: utf-8 -*-
"""
个人实验报告（吕玥 · 后端服务与 OpenClaw 接入）。
沿用《计算机综合项目实践实验报告模板.docx》的格式与样式。
"""
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEMPLATE = "计算机综合项目实践实验报告模板.docx"
OUTPUT = "EmailClaw_个人实验报告_吕玥.docx"

doc = docx.Document(TEMPLATE)

# ---------------- 字体辅助 ----------------
def set_run_font(run, latin="Times New Roman", ea="宋体", size=10.5, bold=False, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin); rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), ea); rfonts.set(qn("w:cs"), latin)
    if size is not None: run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None: run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph(style="Heading 1"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), ea="黑体", size=22, bold=True); return p

def h2(text):
    p = doc.add_paragraph(style="Heading 2")
    set_run_font(p.add_run(text), ea="黑体", size=16, bold=True); return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(text), ea="黑体", size=12, bold=True); return p

def para(text, size=10.5, bold=False, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(3)
    if indent: p.paragraph_format.first_line_indent = Pt(21)
    set_run_font(p.add_run(text), size=size, bold=bold); return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Pt(21 + level * 14)
    p.paragraph_format.space_after = Pt(1)
    set_run_font(p.add_run("• " + text), size=10.5); return p

def code_block(text):
    lines = text.split("\n")
    while lines and lines[0].strip() == "": lines.pop(0)
    while lines and lines[-1].strip() == "": lines.pop()
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Pt(14)
        ppr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
        ppr.append(shd)
        set_run_font(p.add_run(line if line != "" else " "), latin="Consolas", ea="宋体", size=9)

def set_cell_text(cell, text, header=False, align="left", size=9.5):
    cell.text = ""; p = cell.paragraphs[0]
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    set_run_font(p.add_run(str(text)), size=size, bold=header)
    if header:
        tcpr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D9E2F3")
        tcpr.append(shd)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "808080")
        borders.append(el)
    tblPr.append(borders)

def add_table(headers, rows, widths=None, size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_borders(table)
    for i, htext in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], htext, header=True, align="center", size=size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=size)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows: row.cells[i].width = Pt(w)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)
    return table

def caption(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(text), ea="楷体", size=9, color=RGBColor(0x40, 0x40, 0x40))

def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def set_para_text_keepfmt(p, new_text):
    if not p.runs:
        set_run_font(p.add_run(new_text), ea="宋体", size=15); return
    p.runs[0].text = new_text
    for r in p.runs[1:]: r.text = ""

# ---------------- 封面 ----------------
paras = doc.paragraphs
set_para_text_keepfmt(paras[1], "智能邮件管家系统实验报告")
set_para_text_keepfmt(paras[8], "课 程 名 称   ：计算机综合项目实践")
set_para_text_keepfmt(paras[11], "学 生 姓 名   ：吕玥")
set_para_text_keepfmt(paras[17], "二○二六年六月")
set_para_text_keepfmt(paras[42], "本人签名：                     日期：2026年6月6日")

# 删除模板正文（从“摘要”起）
for p in doc.paragraphs[46:]:
    p._element.getparent().remove(p._element)

# ====================================================================
# 摘要
# ====================================================================
page_break()
h1("摘  要")
para("本报告聚焦于小组项目“EmailClaw 智能邮件管家系统”中本人承担的工作部分——后端核心服务与 OpenClaw 大模型接入。EmailClaw 是一套面向多用户的智能邮件处理系统，通过 IMAP 实时监听邮箱、自动完成邮件分类、重要性评分与摘要，并将结果以交互式卡片推送到飞书供用户处理。其中飞书卡片前端与模型 Prompt 调优由另一名同学负责，本人则负责支撑整套系统运转的后端骨架与智能分析的大模型接入。")
para("本人完成的具体工作包括：基于 imap 与 mailparser 实现的多用户 IMAP 长连接管理（连接隔离、断线重连、IDLE 与轮询双触发、并发与跨重启去重）；串联入库、规则、分析、推送的邮件处理流水线；支持“条件—动作”的用户规则引擎；本地关键词分析 Agent 与可选 OpenClaw + DeepSeek 大模型分析的调度与自动回退；基于 Prisma + PostgreSQL 的数据访问与 REST API；以及由 HTTP、业务、数据三层构成的防御性鉴权体系（JWT、bcrypt、AES-256-GCM）。")
para("在 OpenClaw 接入上，本人针对大模型与框架返回的不稳定性做了三处关键工程适配：拆解 OpenClaw 返回的信封结构取出真正的分析 JSON、为每封邮件分配独立会话键以解决并发子进程的会话锁抢占、以及失败重试与代码围栏剥离的容错解析；并保证大模型不可用时自动回退本地规则，使系统始终可用。经端到端联调与越权测试，本人负责的各模块均稳定可靠地完成了预期功能。")
para("关键词：IMAP；多用户连接管理；规则引擎；智能分析 Agent；OpenClaw；DeepSeek；防御性鉴权", indent=False)

# ====================================================================
# 第1章 概述
# ====================================================================
page_break()
h1("第1章 概述")
h2("1.1 选题")
para("电子邮件是工作与学习中最重要的异步通信工具之一，但邮件数量的爆炸式增长带来了严重的信息过载：用户每天面对数十乃至上百封混杂着工作通知、营销推广与垃圾邮件的内容，难以快速识别真正重要、需要立即处理的邮件，也难以从冗长的正文中迅速抓取要点。传统邮件客户端只提供收发与归档，缺乏对内容的语义理解、优先级判断与主动提醒能力。")
para("针对上述痛点，本小组选定并实现了智能邮件管家系统 EmailClaw：作为一个智能中间层实时监听用户邮箱，借助规则与 AI 能力对每封新邮件自动完成分类、重要性评分与摘要，将重要邮件以交互式卡片主动推送给用户，用户在即时通讯工具中即可完成对邮件的常见处理，且操作会真实同步回邮箱。系统进一步支持多用户、自定义规则、邮件检索与防御性鉴权。")
para("本人在该项目中负责后端服务与 OpenClaw 接入，即系统中“感知—分析—执行”链路的服务端实现与大模型分析能力的对接。这部分工作技术密度高、难点集中，是支撑整个系统运转的底座，本报告即围绕这部分内容展开。", indent=True)

h2("1.2 个人分工与职责")
para("项目由两名同学协作完成，分工如下表。本报告聚焦本人承担的后端与 OpenClaw 接入部分，对另一名同学负责的飞书卡片前端与模型调优仅在涉及协作边界时简要提及。")
add_table(
    ["成员", "承担模块", "本报告范围"],
    [
        ["吕玥（本人）", "后端服务与 OpenClaw 接入", "本报告主体：多用户 IMAP、处理流水线、规则引擎、分析 Agent 与 OpenClaw/DeepSeek 接入、数据访问与 REST API、三层鉴权"],
        ["董一鸣", "飞书前端服务与模型参数调优", "协作边界：后端通过 HTTP 向其飞书机器人推送数据、接收其转发的卡片事件"],
    ],
    widths=[70, 130, 305], size=10,
)

# ====================================================================
# 第2章 系统需求分析
# ====================================================================
page_break()
h1("第2章 系统需求分析")
h2("2.1 引言")
para("本章从本人负责模块的角度分析需求。整个系统的功能需求编号为 FR-1 至 FR-7，其中后端服务直接支撑 FR-1、FR-2、FR-3、FR-5、FR-6，并为 FR-4 提供摘要数据来源。需要说明的是，项目在实现中将通知渠道由早期规划的 Telegram 调整为飞书，并由单用户演示形态演进为支持注册登录与凭据加密的多用户系统，下文均以实际实现为准。")
h2("2.2 功能需求")
para("本人负责的后端模块需满足的功能需求如下表所示，其余 UI 呈现类需求由前端承担。")
add_table(
    ["编号", "需求", "后端职责", "状态"],
    [
        ["FR-1", "邮件实时接收与存储（多用户）", "多用户 IMAP 长连接、MIME 解析、按用户隔离入库、去重", "已实现"],
        ["FR-2", "智能分类", "规则引擎 + 本地 Agent + 可选 OpenClaw/DeepSeek 分析与回退", "已实现"],
        ["FR-3", "重要性评分与主动推送", "重要性评分、按用户阈值/偏好做推送决策并调用飞书 bot", "已实现"],
        ["FR-4", "邮件摘要生成", "由分析层生成摘要并持久化（卡片展示由前端完成）", "已实现"],
        ["FR-5", "用户交互与反馈、规则自定义", "真实 IMAP 操作、分类反馈落库、规则 CRUD API", "已实现"],
        ["FR-6", "邮件搜索与过滤", "多条件分页检索 API、按 userId 强制过滤", "已实现"],
        ["FR-7", "每日摘要与统计", "（后续工作）", "未实现"],
    ],
    widths=[40, 150, 250, 50], size=9,
)
caption("表 2-1　本人负责的功能需求")
h2("2.3 非功能需求")
para("后端是系统的安全与稳定底座，非功能需求以安全性、可靠性与性能为重点。")
add_table(
    ["类别", "需求与实现"],
    [
        ["安全性", "密码 bcrypt 哈希；IMAP 授权码 AES-256-GCM 加密入库；用户 API 强制 JWT；webhook 校验共享密钥；查询与写操作强制按 userId 隔离与归属校验"],
        ["可靠性", "IMAP 断线自动重连；大模型分析失败自动回退本地规则；推送失败容错不阻塞主流程；以数据库 notifiedAt 实现跨重启去重"],
        ["性能", "新邮件秒级抓取（IDLE + 30s 兜底轮询）；本地分析毫秒级、大模型分析受网络影响；搜索基于索引分页，千封内 < 1s"],
        ["可维护性", "代码按路由/服务/集成/工具分层；关键日志带 [user=...] 前缀；TypeScript 静态类型；AgentLog 审计"],
    ],
    widths=[60, 450], size=9.5,
)
caption("表 2-2　后端非功能需求")
h2("2.4 技术选型")
add_table(
    ["类别", "技术", "选型说明"],
    [
        ["语言/运行时", "Node.js 20 + TypeScript 5", "异步 I/O 适合邮件与网络密集场景，静态类型利于维护"],
        ["Web 框架", "Express 5", "轻量成熟，便于构建 REST API"],
        ["ORM/数据库", "Prisma 7 + PostgreSQL", "类型安全访问，JSON 字段灵活存偏好与规则"],
        ["邮件", "imap + mailparser", "IMAP 长连接监听 + MIME 解析"],
        ["安全", "bcryptjs / jsonwebtoken / AES-256-GCM", "密码哈希、令牌、凭据加密"],
        ["大模型（可选）", "OpenClaw + DeepSeek(deepseek-chat)", "按需拉起子进程编排，成本低、中文效果好"],
    ],
    widths=[80, 200, 230], size=9.5,
)
caption("表 2-3　后端关键技术选型")
h2("2.5 其他要求")
bullet("运行环境：Node.js 18+、可访问的 PostgreSQL、支持 IMAP 的邮箱及授权码；大模型增强需安装 OpenClaw CLI 并配置 DeepSeek API Key。")
bullet("约束：分类固定 6 类；删除为真实不可恢复删除；状态同步为单向（EmailClaw → 邮箱服务器）。")

# ====================================================================
# 第3章 系统设计
# ====================================================================
page_break()
h1("第3章 系统设计")
h2("3.1 引言")
para("本章详细阐述本人负责的后端设计：先给出后端在整体架构中的位置与数据流，再依次对多用户 IMAP 连接管理、邮件处理流水线、规则引擎、智能分析 Agent 与 OpenClaw/DeepSeek 接入、防御性鉴权、REST API 与数据库进行详细设计。飞书机器人在本章中作为后端的外部协作边界出现，不展开其卡片构建细节。")

h2("3.2 后端总体设计与数据流")
para("系统采用双服务协作架构，本人负责的主后端服务（端口 3000）处于核心位置，负责邮件抓取解析、智能分析、数据持久化、真实邮箱操作与 REST API；前端飞书机器人服务（端口 3001）通过 HTTP 与后端互调。后端处理一封新邮件的数据流如下图所示。")
code_block(r"""
 ┌────────────┐  IMAP长连接   ┌──────────────────────────────────────────┐
 │ 用户邮箱    │◀────────────▶│            主后端服务 (:3000) [本人]       │
 └────────────┘  抓取/操作     │  UserMailbox ── ImapManager                │
                              │        │ onIncomingEmail                    │
   ┌────────────┐  Prisma     │        ▼  processIncomingEmail:            │
   │ PostgreSQL │◀───────────▶│  入库→去重→规则引擎→分析(本地/OpenClaw)→   │
   └────────────┘             │  推送决策→标记已处理                        │
                              └─────┬───────────────────────▲──────────────┘
                       POST /api/notify-email      POST /api/feishu/webhook
                                    ▼ (推送)         │ (回调, X-Bot-Secret)
                              ┌──────────────────────────────────────────┐
                              │     飞书机器人服务 (:3001) [协作边界]      │
                              └──────────────────────────────────────────┘
""")
caption("图 3-1　后端在系统中的位置与数据流")

h2("3.3 后端详细设计")

h3("3.3.1 多用户 IMAP 连接管理")
para("为支持多用户，将 IMAP 连接抽象为两层：UserMailbox 封装单个用户的长连接与邮件操作，ImapManager 作为多用户连接注册表统一调度。每个绑定邮箱的用户对应一个隔离的连接实例，日志带 [user=...] 前缀。")
bullet("连接与重连：建立 TLS 长连接并开启 keepalive；监听 error/end，断开后延迟 5 秒自动重连。")
bullet("新邮件感知：连接就绪后打开 INBOX，先全量扫描未读，再监听 mail 事件；考虑到 QQ 等邮箱 IDLE 推送不可靠，额外加 30 秒兜底轮询。")
bullet("去重设计：服务器侧写入自定义关键字 CLAWED；内存维护 claimedUids 集合与 scanning 重入锁挡住标记写入前的并发重复；最终以数据库 notifiedAt 作为跨重启去重的权威依据（部分邮箱不持久化自定义关键字）。")
bullet("配置推断：resolveImapConfig 按邮箱域名推断 QQ/163/Gmail/Outlook 等服务商的 IMAP 主机端口；归档目录不存在时自动创建。")
para("ImapManager 对外提供 startForAllBoundUsers / startForUser / stopForUser / getMailboxForEmail；其中 getMailboxForEmail 按 emailId 路由到对应连接时会再次做归属校验，是数据层防御的一环。")

h3("3.3.2 邮件处理流水线")
para("新邮件统一进入 processIncomingEmail 流水线，串联入库、去重、规则、分析、推送决策与标记六个步骤，核心逻辑如下。")
code_block(r"""
processIncomingEmail(userId, email):
  saved = upsertEmail(userId, email)            # 1. 入库(userId+uid 唯一)
  if saved.notifiedAt: return                    # 2. 已处理 → 跳过(跨重启去重)
  ruleHit = ruleEngine.evaluate(userId, email)   # 3. 先跑规则
  if ruleHit:
      analysis = ruleHit.result                  #    命中 → 跳过 Agent
      applyRuleSideEffects(ruleHit.sideEffects)  #    执行 mark_read/archive/delete
  else:
      analysis = agentService.analyzeEmail()     # 4. 未命中 → Agent 分析
  isImportant = analysis.importance >= (threshold ?? 7)
  if pushAllEmails or isImportant:               # 5. 按偏好推送飞书卡片
      feishuService.pushEmailCard(..., isImportant)
  markNotified(saved.id)                          # 6. 标记已处理
""")
caption("图 3-2　邮件处理流水线伪代码")

h3("3.3.3 规则引擎")
para("规则引擎在 Agent 之前运行，命中即跳过 Agent，让用户对确定性场景拥有更高优先级、更可控的处理。规则存于 Rule 表，conditions/actions 为 JSON，约定如下。")
code_block(r"""
conditions: [                       // 多条件 AND
  { field: from|to|subject|body,
    operator: contains|equals|startsWith|endsWith|regex,
    value: "...", caseSensitive?: false } ]
actions: { category: 6类之一(必填), importance?: 0-10,
           summary?: "自定义", sideEffects?: [mark_read|archive|delete] }
""")
para("引擎按 priority 降序遍历用户启用规则，要求其全部条件命中（every）；首个命中者合成分析结果（置信度记为 1、模型标记 rule-engine-v1）并携带副作用动作返回。另提供 validateDefinition 在创建/更新规则时校验字段合法性与正则可编译性。")

h3("3.3.4 智能分析 Agent 与 OpenClaw + DeepSeek 接入")
para("AgentService 是分析层统一入口：analyzeEmailDraft 优先尝试 OpenClaw（若启用），失败回退本地规则 Agent；analyzeEmail 在此基础上持久化结果与 AgentLog。本地规则 Agent 由 ClassificationSkill 实现，包含分类（归一化文本后按 spam→work→shopping→marketing→personal 顺序关键词匹配）、重要性评分（基础 3 分，命中紧急/截止/会议等 +4，含附件 +1，上限 10）与摘要三项能力，不依赖任何外部服务，作为大模型的兜底，保证系统始终可用。")
para("当配置 OPENCLAW_ENABLED=true 时，OpenClawClient 为每封邮件临时拉起一个 OpenClaw 子进程，以本机内嵌（--local）模式调用 DeepSeek 完成分析。这是本人工作中的难点，命令形如：")
code_block(r"""
openclaw agent --agent email-claw --session-key <每封唯一> \
        --message <结构化邮件提示词> --json --local
""")
para("由于大模型与框架返回的不稳定性，OpenClawClient 做了三处关键适配（这三处是“能跑通”与“看似配好实则全是垃圾分类”的分水岭）：")
bullet("拆信封解析：OpenClaw 的 --json 返回 {payloads, meta} 信封，真正的分析 JSON 套在 payloads[0].text 字符串里，需先拆出再解析，否则会静默拿到默认值，每封都退化为 other / 重要性 3。")
bullet("独立会话键：后端并发拉起多个 --local 子进程会抢同一会话文件锁，触发 EmbeddedAttemptSessionTakeoverError；为每封邮件生成唯一 session-key 实现进程间隔离。")
bullet("容错与重试：DeepSeek 偶发把 JSON 包进 ```json``` 围栏或输出非法 JSON，客户端剥除围栏、用正则兜底提取，并在首次失败后换 session-key 重试一次；两次都失败才抛出，由上层回退本地规则。")
para("分析结果经 normalizeResult 归一化：分类越界归为 other、置信度裁剪到 [0,1]、重要性裁剪到 [0,10]、缺省字段填默认值，确保下游拿到结构稳定的结果。整体形成“规则引擎 → 本地 Agent / OpenClaw+DeepSeek → 失败自动回退”的分层智能策略，在成本、稳定与效果间取得务实平衡。")

h3("3.3.5 防御性鉴权与数据安全")
para("针对多账号邮件场景的安全需求，本人设计了 HTTP、业务、数据三层防御性鉴权，层层校验确保用户无法越权访问或操作他人邮件。")
add_table(
    ["层级", "位置", "校验内容"],
    [
        ["HTTP 层", "requireAuth / requireBotSecret", "用户 API 强制 JWT；webhook 必须携带正确 X-Bot-Secret"],
        ["业务层", "feishuRoutes", "先将飞书 openId 解析为 userId，未绑定者直接 403"],
        ["数据层", "assertEmailOwnership / where userId", "写操作前校验 emailId 归属；规则增删改强制带 userId"],
    ],
    widths=[55, 200, 255], size=9.5,
)
caption("表 3-1　防御性鉴权三层防线")
para("配套安全工具：crypto 模块以 AES-256-GCM 加解密 IMAP 授权码（密钥由 ENCRYPTION_KEY 经 SHA-256 派生，密文为 iv‖密文‖认证标签 的 base64）；auth 模块以 bcrypt（12 轮）哈希密码、以 jsonwebtoken 签发校验 7 天有效期 JWT。引导用户机制（bootstrap）在启动时按 .env 自动注册/补齐引导用户并加密绑定邮箱，兼容既有单用户配置。")

h3("3.3.6 REST API 设计")
add_table(
    ["方法", "路径", "说明"],
    [
        ["POST", "/api/auth/register · /login", "注册/登录，返回 JWT"],
        ["GET/PATCH", "/api/users/me · /preferences · /feishu", "用户信息、偏好(阈值/全量推送)、绑定飞书"],
        ["POST/DELETE", "/api/users/me/mailbox", "绑定(加密入库并启动 IMAP)/解绑邮箱"],
        ["GET", "/api/emails · /:id", "多条件分页搜索 / 单封详情(校验归属)"],
        ["GET/POST/PATCH/DELETE", "/api/rules", "规则增删改查(按优先级降序、userId 隔离)"],
        ["POST", "/api/feishu/webhook", "飞书卡片事件回调(需 X-Bot-Secret)"],
    ],
    widths=[110, 200, 200], size=9,
)
caption("表 3-2　后端 REST API")

h2("3.4 数据库设计")
para("系统使用 PostgreSQL，经 Prisma 建模，共 6 张表：User、Email、Classification、Contact、Rule、AgentLog。本人负责其建模与数据访问层（databaseService）。核心表与关系如下。")
code_block(r"""
 User(1) ─< Email(N) ─< Classification(1对1)
   ├─< Rule(N)   ├─< Contact(N)   └─< AgentLog(N)
 Email 以 (userId, uid) 唯一；索引 (userId, receivedAt)、(category)
""")
caption("图 3-3　数据库 ER 关系")
add_table(
    ["表", "关键字段", "说明"],
    [
        ["User", "email, password, imapPassword(加密), feishuUserId, preferences(JSON)", "用户与凭据、偏好"],
        ["Email", "userId, uid, from/subject/body, category, importance, summary, isRead/Archived/Deleted, notifiedAt", "邮件主体与状态，notifiedAt 用于去重"],
        ["Classification", "emailId(唯一), category, confidence, reasoning, model, feedback", "分类结果与用户反馈"],
        ["Rule", "userId, conditions(JSON), actions(JSON), priority, isEnabled", "用户自定义规则"],
        ["AgentLog", "type, status, input/output(JSON), model, duration", "分析执行审计日志"],
    ],
    widths=[75, 270, 165], size=9,
)
caption("表 3-3　主要数据表说明")

# ====================================================================
# 第4章 系统实现
# ====================================================================
page_break()
h1("第4章 系统实现")
h2("4.1 引言")
para("本章介绍本人负责模块的实现情况。系统在 Linux/WSL2 + Node.js 20 环境开发，后端用 ts-node-dev 热重载启动，TypeScript 编译验证通过。本人负责的后端源码结构如下。")
code_block(r"""
src/server.ts                        后端入口(:3000)，挂载路由 + bootstrap
src/backend/
├── api/routes/                      认证/用户/邮件/规则/飞书回调 路由
├── middleware/authMiddleware.ts     requireAuth / requireBotSecret
├── services/
│   ├── emailService.ts              邮件处理流水线 facade
│   ├── userMailbox.ts               单用户 IMAP 长连接封装
│   ├── imapManager.ts               多用户 IMAP 连接注册表
│   ├── ruleEngine.ts                规则匹配引擎
│   ├── agentService.ts              本地 Agent + OpenClaw 调度
│   ├── openClawClient.ts            OpenClaw/DeepSeek 子进程适配 ★
│   ├── databaseService.ts           Prisma 数据访问
│   └── bootstrap.ts                 启动自检 / 引导用户
├── integrations/feishu/feishuService.ts   推送 + 回调分发
└── utils/{auth.ts, crypto.ts}       JWT/bcrypt、AES-256-GCM
src/agents/                          本地规则 Agent、Skill 与工具
prisma/schema.prisma                 6 张表数据模型 + 迁移
""")
caption("图 4-1　本人负责的后端源码结构")

h2("4.2 关键实现与运行")
h3("REST API 调用示例")
code_block(r"""
# 注册
curl -X POST :3000/api/auth/register -H 'Content-Type: application/json' \
  -d '{"email":"u@example.com","password":"pass1234"}'   # → { token, user }
# 绑定 IMAP(需令牌，密码加密入库并后台启动连接)
curl -X POST :3000/api/users/me/mailbox -H "Authorization: Bearer $TOKEN" \
  -d '{"imapHost":"imap.qq.com","imapUser":"x@qq.com","imapPassword":"授权码"}'
# 多条件搜索
curl ":3000/api/emails?category=work&importanceMin=7&q=会议&page=1" \
  -H "Authorization: Bearer $TOKEN"   # → { items, total, page, ... }
""")
caption("图 4-2　REST API 调用示例")
h3("OpenClaw + DeepSeek 接入步骤")
para("本人编写了 Linux 与 Windows 两版接入文档。核心步骤为：全局安装 openclaw → 配置 DeepSeek API Key 与默认模型 deepseek/deepseek-chat → 创建 email-claw agent → 单独自测连通与分析 JSON → 在后端 .env 开启 OPENCLAW_ENABLED 等四项配置。判定增强生效的标准是：分类不再清一色 other、重要性不再清一色 3，日志无 TakeoverError 与解析失败。")
code_block(r"""
openclaw config set models.providers.deepseek.apiKey "sk-..."
openclaw config set agents.defaults.model "deepseek/deepseek-chat"
openclaw agents add email-claw --model deepseek/deepseek-chat --non-interactive
# 后端 .env: OPENCLAW_ENABLED=true / OPENCLAW_COMMAND=openclaw /
#            OPENCLAW_AGENT_ID=email-claw / OPENCLAW_TIMEOUT_MS=60000
""")
caption("图 4-3　OpenClaw + DeepSeek 接入关键配置")

h2("4.3 测试计划")
para("本人针对负责模块设计了功能测试与安全（越权）测试。测试环境：Linux/WSL2 + Node 20 + 本地 PostgreSQL；邮箱用 QQ 邮箱（IMAP 授权码）；大模型用 OpenClaw + DeepSeek。主要用例如下。")
add_table(
    ["编号", "测试项", "预期结果"],
    [
        ["TC-01", "多用户邮件抓取入库", "秒级抓取解析入库，按用户隔离，不重复"],
        ["TC-02", "本地规则分类", "正确分入 work/marketing/other 等"],
        ["TC-03", "OpenClaw+DeepSeek 分析", "分类非清一色 other、重要性非清一色 3"],
        ["TC-04", "失败回退", "禁用大模型后自动回退本地规则，仍正常收邮件"],
        ["TC-05", "规则命中与副作用", "命中规则、跳过 Agent、执行真实 IMAP 副作用"],
        ["TC-06", "重要性推送决策", "达阈值标记重要推送，未达按偏好处理"],
        ["TC-07", "邮件搜索", "多条件组合查询结果正确、分页正常"],
        ["TC-08", "认证与越权防御", "无/错令牌被拦截；越权读写被三层防线拦截"],
    ],
    widths=[45, 150, 315], size=9,
)
caption("表 4-1　测试用例")

h2("4.4 测试报告")
para("本人对负责模块进行了真实启动与端到端联调，结果如下。")
bullet("功能联调：引导用户自动建成并加密绑定邮箱；IMAP 鉴权成功并自动创建 Archive 目录；未读邮件被抓取、解析、入库并完成分析；登录获取 JWT 后 /api/users/me、邮件搜索、规则 CRUD 全部通过；规则命中可跳过 Agent 并执行真实副作用。")
bullet("大模型验证：按接入文档配置后，分类不再退化为清一色 other、重要性不再清一色 3，日志无 TakeoverError 与解析失败；将 OPENCLAW_ENABLED 置为 false 重启后正确回退本地规则且仍能正常收发。")
bullet("安全测试：攻击者用自身 JWT 读他人邮件返回“无权操作此邮件”，列他人邮件返回空集；无 X-Bot-Secret 调 webhook 返回 401；未绑定飞书的 openId 调 webhook 返回 403。三层防线全部命中。")
para("联调中本人定位并根治了“同一封邮件重复发卡”的缺陷，其跟踪情况如下。")
add_table(
    ["缺陷", "根因", "修复"],
    [
        ["重复发卡", "IMAP 轮询并发 + CLAWED 标记写入存在时间窗，跨重启又会重扫",
         "新增内存级 claimedUids 锁与 scanning 重入锁，并以数据库 notifiedAt 字段做权威去重"],
    ],
    widths=[55, 230, 225], size=9,
)
caption("表 4-2　缺陷跟踪与修复")
para("综合结果，本人负责的后端已稳定打通“IMAP 抓取 → MIME 解析 → 入库 → 规则/Agent 分析 → 推送决策 → 真实邮箱与数据库同步”的服务端全链路，覆盖了 FR-1、FR-2、FR-3、FR-5、FR-6 的后端能力。")

# ====================================================================
# 第5章 总结
# ====================================================================
page_break()
h1("第5章 总结")
para("在本次综合项目实践中，本人负责 EmailClaw 智能邮件管家系统的后端核心服务与 OpenClaw 大模型接入，完成了支撑整套系统运转的服务端实现。具体成果包括：多用户 IMAP 长连接管理（连接隔离、断线重连、IDLE 与轮询双触发、并发与跨重启去重）、串联入库与分析与推送的邮件处理流水线、用户规则引擎、本地分析 Agent 与 OpenClaw+DeepSeek 大模型分析的调度与自动回退、基于 Prisma+PostgreSQL 的数据访问与 REST API，以及 HTTP/业务/数据三层防御性鉴权。")
para("本人工作的难点与收获主要集中在两方面：其一是多用户场景下的 IMAP 并发与去重，IMAP IDLE 推送在不同邮箱服务商上可靠性不一，叠加轮询兜底后又带来并发重复处理问题，最终通过内存锁与数据库标记的双重机制才得以根治，让我深刻体会到分布式/并发场景下“恰好一次”处理的工程难度；其二是 OpenClaw + DeepSeek 的接入，大模型与框架返回结构的不稳定决定了直接调用无法跑通，通过拆信封、独立会话键、容错重试三处适配，并辅以失败自动回退，才把链路成功率拉到可用水平，这让我认识到 AI 能力工程化落地中“稳定性兜底”与“效果”同样重要。")
para("不足之处在于：当前重要性评分与分类仍以规则与关键词为主、大模型反馈学习闭环尚未打通，FR-7 每日摘要与统计、自动化测试覆盖也有待补充。后续可在 Agent 反馈学习、定时统计任务与生产级稳定性打磨方向继续完善。")
para("通过这部分工作，本人完整经历了从需求分析、模块设计、编码实现到测试联调与缺陷修复的全过程，在网络协议、数据库、异步并发、第三方平台对接、AI Agent 编排与系统安全等方面都得到了切实锻炼，对工程化落地一个贴近真实形态的系统有了更深入的理解。")

doc.save(OUTPUT)
print("个人报告已生成：", OUTPUT)
