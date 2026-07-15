# -*- coding: utf-8 -*-
"""
把 v1 报告中 4 张 ASCII 示意图（用例图/功能模块图/架构数据流图/ER 图）替换为 PNG 图片。
仅改动这 4 个图块，其余内容（含伪代码、JSON、源码目录、curl 示例、用户的所有手改）保持不变。
"""
import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = "docs/reports/EmailClaw_实验报告v1.docx"
OUT = "docs/reports/EmailClaw_实验报告v1_配图版.docx"

# 按 caption 关键字映射到图片
MAP = [
    ("用例图", "report_images/fig_usecase.png", 14.6),
    ("功能模块图", "report_images/fig_module.png", 15.0),
    ("架构", "report_images/fig_arch.png", 13.2),
    ("ER", "report_images/fig_er.png", 14.6),
]

doc = docx.Document(SRC)

def is_shaded(p):
    ppr = p._p.find(qn('w:pPr'))
    if ppr is None: return False
    shd = ppr.find(qn('w:shd'))
    return shd is not None and shd.get(qn('w:fill')) == 'F2F2F2'

# 1) 检测所有底纹代码块 → (起段, 止段, caption)
paras = doc.paragraphs
blocks = []
i = 0
while i < len(paras):
    if is_shaded(paras[i]):
        s = i
        while i < len(paras) and is_shaded(paras[i]):
            i += 1
        e = i - 1
        cap = ""
        for j in range(e + 1, min(e + 4, len(paras))):
            if paras[j].text.strip():
                cap = paras[j].text.strip(); break
        blocks.append((s, e, cap))
    else:
        i += 1

# 2) 选出 4 个目标块（按 caption 关键字），并确定图片
targets = []  # (start, end, img, width_cm, cap)
for (s, e, cap) in blocks:
    for kw, img, w in MAP:
        if kw in cap and ("图" in cap or kw == "ER"):  # 仅匹配“图 x-x”标题块
            targets.append((s, e, img, w, cap)); break

print("将替换以下图块：")
for (s, e, img, w, cap) in targets:
    print(f"  段{s}..{e} ({e-s+1}行) → {img}  | {cap[:40]}")

# 3) 倒序替换，保证索引有效；首段作为图片占位，删除其余段
para_objs = doc.paragraphs  # 固定引用列表
for (s, e, img, w, cap) in sorted(targets, key=lambda t: -t[0]):
    holder = para_objs[s]
    # 清空首段内容 + 去底纹 + 去缩进 + 居中
    holder.clear()
    ppr = holder._p.find(qn('w:pPr'))
    if ppr is not None:
        for tag in ('w:shd', 'w:ind'):
            el = ppr.find(qn(tag))
            if el is not None:
                ppr.remove(el)
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.paragraph_format.line_spacing = 1.0
    holder.add_run().add_picture(img, width=Cm(w))
    # 删除该块剩余段落
    for p in para_objs[s + 1:e + 1]:
        p._element.getparent().remove(p._element)

doc.save(OUT)
print("\n已生成：", OUT)
