# -*- coding: utf-8 -*-
"""
个人实验报告（董一鸣 · 飞书前端服务与模型参数调优）。
沿用《计算机综合项目实践实验报告模板.docx》的格式与样式。
"""
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEMPLATE = "计算机综合项目实践实验报告模板.docx"
OUTPUT = "EmailClaw_个人实验报告_董一鸣.docx"
doc = docx.Document(TEMPLATE)

def set_run_font(run, latin="Times New Roman", ea="宋体", size=10.5, bold=False, color=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin); rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), ea); rfonts.set(qn("w:cs"), latin)
    if size is not None: run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None: run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph(style="Heading 1"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), ea="黑体", size=22, bold=True); return p
def h2(text):
    p = doc.add_paragraph(style="Heading 2")
    set_run_font(p.add_run(text), ea="黑体", size=16, bold=True); return p
def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(text), ea="黑体", size=12, bold=True); return p
def para(text, size=10.5, bold=False, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(3)
    if indent: p.paragraph_format.first_line_indent = Pt(21)
    set_run_font(p.add_run(text), size=size, bold=bold); return p
def bullet(text, level=0):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Pt(21 + level * 14)
    p.paragraph_format.space_after = Pt(1)
    set_run_font(p.add_run("• " + text), size=10.5); return p
def code_block(text):
    lines = text.split("\n")
    while lines and lines[0].strip() == "": lines.pop(0)
    while lines and lines[-1].strip() == "": lines.pop()
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Pt(14)
        ppr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2F2")
        ppr.append(shd)
        set_run_font(p.add_run(line if line != "" else " "), latin="Consolas", ea="宋体", size=9)
def set_cell_text(cell, text, header=False, align="left", size=9.5):
    cell.text = ""; p = cell.paragraphs[0]
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    set_run_font(p.add_run(str(text)), size=size, bold=header)
    if header:
        tcpr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D9E2F3")
        tcpr.append(shd)
def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "808080")
        borders.append(el)
    tblPr.append(borders)
def add_table(headers, rows, widths=None, size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; set_table_borders(table)
    for i, htext in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], htext, header=True, align="center", size=size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=size)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows: row.cells[i].width = Pt(w)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)
    return table
def caption(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(text), ea="楷体", size=9, color=RGBColor(0x40, 0x40, 0x40))
def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
def set_para_text_keepfmt(p, new_text):
    if not p.runs:
        set_run_font(p.add_run(new_text), ea="宋体", size=15); return
    p.runs[0].text = new_text
    for r in p.runs[1:]: r.text = ""

# ---------------- 封面 ----------------
paras = doc.paragraphs
set_para_text_keepfmt(paras[1], "智能邮件管家系统实验报告")
set_para_text_keepfmt(paras[8], "课 程 名 称   ：计算机综合项目实践")
set_para_text_keepfmt(paras[11], "学 生 姓 名   ：董一鸣")
set_para_text_keepfmt(paras[17], "二○二六年六月")
set_para_text_keepfmt(paras[42], "本人签名：                     日期：2026年6月6日")
for p in doc.paragraphs[46:]:
    p._element.getparent().remove(p._element)

# ====================================================================
# 摘要
# ====================================================================
page_break()
h1("摘  要")
para("本报告聚焦于小组项目“EmailClaw 智能邮件管家系统”中本人承担的工作部分——飞书机器人前端服务与大模型参数调优。EmailClaw 是一套面向多用户的智能邮件处理系统，通过 IMAP 实时监听邮箱、自动完成邮件分类、重要性评分与摘要，并将结果以交互式卡片推送到飞书供用户处理。其中后端核心服务与 OpenClaw 接入由另一名同学负责，本人则负责面向用户的飞书交互入口与决定分析效果的大模型 Prompt 与参数调优。")
para("本人完成的具体工作包括：基于飞书开放平台 Node SDK 搭建飞书机器人服务，通过 WebSocket 长连接接收卡片按钮事件并提供 HTTP 接口接收后端推送；设计并实现邮件通知卡片、邮件详情卡片与分类选择卡片三类交互式卡片，包括按分类着色、重要邮件高亮、星级重要性、状态展示与三行共八个操作按钮；针对飞书卡片回调的 3 秒超时限制，设计了“先返回 toast、再异步刷新卡片”的交互模式，打通了卡片状态刷新与详情展示的完整闭环。")
para("在模型参数调优方面，本人负责 OpenClaw + DeepSeek 大模型分析的 Prompt 设计与参数调优：通过强约束的结构化提示词、明确的输出 JSON 字段规范、正文截断与转义要求等手段，配合模型选型与超时参数的调整，显著降低了分类退化为“其他”、重要性退化为默认值与输出非法 JSON 的概率，使大模型在分类、重要性与摘要三项任务上的输出质量与稳定性达到可用水平。经端到端联调，本人负责的前端卡片交互与模型分析效果均达到了预期。")
para("关键词：飞书机器人；交互式卡片；卡片刷新；WebSocket；Prompt 设计；DeepSeek；参数调优", indent=False)

# ====================================================================
# 第1章 概述
# ====================================================================
page_break()
h1("第1章 概述")
h2("1.1 选题")
para("电子邮件是工作与学习中最重要的异步通信工具之一，但邮件数量的爆炸式增长带来了严重的信息过载：用户每天面对数十乃至上百封混杂着工作通知、营销推广与垃圾邮件的内容，难以快速识别真正重要的邮件，也难以从冗长正文中迅速抓取要点。传统邮件客户端只提供收发与归档，缺乏对内容的语义理解、优先级判断与主动提醒能力。")
para("针对上述痛点，本小组选定并实现了智能邮件管家系统 EmailClaw：作为智能中间层实时监听邮箱，借助规则与 AI 能力对每封新邮件自动完成分类、重要性评分与摘要，将重要邮件以交互式卡片主动推送给用户，用户在常用的即时通讯工具（飞书）中即可完成对邮件的常见处理，且操作会真实同步回邮箱。")
para("本人在该项目中负责飞书机器人前端服务与大模型参数调优。前端是用户与系统交互的唯一入口，直接决定产品体验；模型调优则直接决定分类、重要性与摘要的分析质量。这两部分一个面向“好不好用”，一个面向“准不准”，是系统价值能否被用户感知到的关键，本报告即围绕这部分内容展开。", indent=True)

h2("1.2 个人分工与职责")
para("项目由两名同学协作完成，分工如下表。本报告聚焦本人承担的飞书前端与模型调优部分，对另一名同学负责的后端服务与 OpenClaw 工程接入仅在涉及协作边界时简要提及。")
add_table(
    ["成员", "承担模块", "本报告范围"],
    [
        ["董一鸣（本人）", "飞书前端服务与模型参数调优", "本报告主体：飞书机器人服务、三类交互式卡片、卡片刷新交互闭环、大模型 Prompt 设计与参数调优、效果评估"],
        ["吕玥", "后端服务与 OpenClaw 接入", "协作边界：后端推送邮件分析数据、接收并执行卡片按钮事件、提供 OpenClaw 子进程工程接入"],
    ],
    widths=[80, 130, 295], size=10,
)

# ====================================================================
# 第2章 系统需求分析
# ====================================================================
page_break()
h1("第2章 系统需求分析")
h2("2.1 引言")
para("本章从本人负责模块的角度分析需求。整个系统的功能需求编号为 FR-1 至 FR-7，其中飞书前端服务直接支撑 FR-3（重要邮件展示与高亮）、FR-4（摘要展示）、FR-5（用户交互与反馈）；模型参数调优则直接服务于 FR-2（分类质量）、FR-3（重要性合理性）、FR-4（摘要凝练度）。需要说明的是，项目在实现中将交互渠道由早期规划的 Telegram 调整为卡片交互能力更强的飞书（Lark），下文均以实际实现为准。")
h2("2.2 功能需求")
para("本人负责的飞书前端与模型调优需满足的功能需求如下表所示，其余服务端逻辑由后端承担。")
add_table(
    ["编号", "需求", "本人职责", "状态"],
    [
        ["FR-2", "智能分类（质量）", "通过 Prompt 与参数调优提升大模型分类准确性、降低退化", "已实现"],
        ["FR-3", "重要性评分与主动推送", "卡片重要邮件红色高亮+🔥，星级展示重要性", "已实现"],
        ["FR-4", "邮件摘要生成（展示）", "卡片展示摘要；调优使摘要更凝练", "已实现"],
        ["FR-5", "用户交互与反馈", "三类卡片、8 个操作按钮、回传事件、卡片刷新与详情", "已实现"],
    ],
    widths=[40, 150, 270, 50], size=9,
)
caption("表 2-1　本人负责的功能需求")
h2("2.3 非功能需求")
para("飞书前端是用户直接感知的界面，非功能需求以用户体验与交互可靠性为重点；模型调优则以输出稳定性为重点。")
add_table(
    ["类别", "需求与实现"],
    [
        ["用户体验", "卡片格式美观、信息层次清晰；重要邮件高亮；删除操作二次确认；操作后即时反馈并刷新卡片状态"],
        ["交互可靠性", "针对飞书卡片回调 ~3 秒超时，采用“先 toast 再异步刷新”模式，避免红色超时错误；操作失败有明确文字提示，不静默吞错"],
        ["模型输出稳定性", "通过强约束 Prompt 与归一化，降低分类退化、重要性退化与非法 JSON 概率，使有效解析率接近 100%"],
        ["可用性", "提供 mock 卡片接口，无需后端即可独立调试卡片样式与交互"],
    ],
    widths=[70, 440], size=9.5,
)
caption("表 2-2　前端与模型非功能需求")
h2("2.4 技术选型")
add_table(
    ["类别", "技术", "选型说明"],
    [
        ["运行时", "Node.js 20", "与后端一致，便于协作"],
        ["飞书对接", "飞书开放平台 Node SDK(@larksuiteoapi/node-sdk)", "官方 SDK，封装消息发送、卡片更新与事件分发"],
        ["事件通道", "WebSocket 长连接(WSClient)", "免公网回调地址，内网联调便捷，实时接收卡片事件"],
        ["HTTP 框架", "Express", "提供 /api/notify-email 等接收后端推送的接口"],
        ["大模型", "OpenClaw + DeepSeek(deepseek-chat)", "成本低、中文效果好，作为分类/重要性/摘要的增强引擎"],
    ],
    widths=[60, 230, 220], size=9.5,
)
caption("表 2-3　前端与模型关键技术选型")
h2("2.5 其他要求")
bullet("运行环境：Node.js 18+；飞书企业自建应用（开启机器人能力、消息发送权限与卡片事件订阅），配置 APP_ID/APP_SECRET 与默认接收者 DEFAULT_OPEN_ID；与后端共享 FEISHU_BOT_SHARED_SECRET。")
bullet("约束：卡片不适合展示超长正文，详情卡片仅展示正文前若干字符；分类枚举固定 6 类；按钮 action 命名须与后端约定一致。")

# ====================================================================
# 第3章 系统设计
# ====================================================================
page_break()
h1("第3章 系统设计")
h2("3.1 引言")
para("本章详细阐述本人负责的飞书前端与模型调优设计：先给出飞书机器人服务在整体架构中的位置与交互流程，再对三类交互式卡片的 UI、卡片刷新交互闭环进行详细设计，最后阐述大模型 Prompt 设计与参数调优方案。后端服务在本章中作为外部协作边界出现。")

h2("3.2 飞书前端总体设计与交互流程")
para("飞书机器人服务（端口 3001）是用户交互入口，本身不直接访问邮箱与数据库，只负责展示后端推送的分析结果、提供操作按钮、并将用户点击事件回传后端。它通过两条链路与后端协作：推送链路（后端 → bot → 飞书用户）与回调链路（飞书用户 → bot → 后端），如下图所示。")
code_block(r"""
  后端(:3000)                        飞书机器人服务(:3001) [本人]            飞书用户
     │  POST /api/notify-email           │                                     │
     ├──────────────────────────────────▶│ buildEmailCard()                    │
     │  (邮件分析数据)                    │ feishuClient.sendCard ──────────────▶│ 收到卡片
     │                                   │                                     │
     │                                   │ WSClient 监听 card.action.trigger ◀─┤ 点击按钮
     │                                   │ ① 先同步返回 toast("处理中...")     │ (停止3秒倒计时)
     │  POST /api/feishu/webhook         │ ② 后台 forwardToBackend(带密钥)     │
     │◀──────────────────────────────────┤                                     │
     │  执行业务并回传最新邮件状态        │                                     │
     ├──────────────────────────────────▶│ ③ updateCard 刷新原卡片 ────────────▶│ 卡片状态更新
     │                                   │   或 sendCard 新发详情卡片          │
""")
caption("图 3-1　飞书机器人服务的推送链路与回调链路")
para("机器人服务通过 feishuClient 封装飞书 OpenAPI，提供 sendCard（发送交互式卡片）、updateCard（patch 已发卡片）、sendText（文本提示）与 getMessage 等方法；并提供 /api/notify-email、/api/mock-card、/api/send-text 等 HTTP 接口，其中 mock-card 用于无后端独立调试卡片。")

h2("3.3 飞书卡片 UI 设计")
para("系统的用户界面以三类交互式卡片为载体：邮件通知卡片、邮件详情卡片、分类选择卡片。其中邮件通知卡片是核心，布局如下图所示。")
code_block(r"""
┌──────────────────────────────────────────────┐
│ 🔥 邮件主题            （重要邮件红色卡头+🔥）  │ ← header 按分类着色
├──────────────────────────────────────────────┤
│ 发件人：xxx           收件人：yyy              │
│ 时间：2026-06-06 10:00                          │
│ ──────────────────────────────────────────    │
│ 📂 分类：工作 置信度90%  ⭐重要性 ★★★★★★★★(8/10)│
│ ──────────────────────────────────────────    │
│ 📝 摘要：……    💡 分类理由：……                │
│ 状态：📬未读 | 📂收件箱                          │
│ ──────────────────────────────────────────    │
│ [📖标为已读][⭐标为重点][📦归档]               │
│ [✅分类正确][❌分类错误][🔍查看详情]           │
│ [🗑删除]   [🔄重新分析]                         │
└──────────────────────────────────────────────┘
""")
caption("图 3-2　飞书邮件通知卡片布局")
para("卡片设计上的若干要点：按邮件分类映射卡头颜色以便快速区分类型，当邮件被判定为重要时统一覆盖为红色并加“🔥”前缀；重要性以 0~10 的星级（★/☆）直观呈现；状态行显示已读/未读与归档状态；删除按钮带二次确认弹窗，提示该操作真实不可恢复；为避免飞书 Markdown 渲染异常，对所有用户文本做 escapeMd 转义，时间统一格式化为东八区可读字符串。分类与颜色映射、按钮设计分别如下两表。")
add_table(
    ["分类", "work", "personal", "shopping", "marketing", "spam", "other"],
    [["中文", "工作", "个人", "购物", "营销", "垃圾", "其他"],
     ["卡头色", "蓝", "绿", "黄", "紫", "红", "灰"]],
    size=9.5,
)
caption("表 3-1　分类与卡片卡头颜色映射")
add_table(
    ["按钮", "action", "回传内容", "交互"],
    [
        ["标为已读", "mark_read", "emailId", "异步刷新卡片状态"],
        ["标为重点", "mark_important", "emailId", "异步刷新"],
        ["归档", "archive", "emailId", "异步刷新"],
        ["删除", "delete", "emailId", "二次确认 → 异步刷新（按钮区消失）"],
        ["分类正确", "feedback_correct", "emailId", "异步刷新"],
        ["分类错误", "feedback_wrong", "emailId(+expectedCategory)", "弹出分类选择卡片 → 选择后回传"],
        ["查看详情", "view_detail", "emailId", "新发一张详情卡片，保留原卡片"],
        ["重新分析", "reanalyze", "emailId", "异步刷新"],
    ],
    widths=[60, 95, 175, 180], size=9,
)
caption("表 3-2　卡片按钮设计")
para("分类选择卡片在用户点击“分类错误”时同步弹出，列出 work/personal/shopping/marketing/spam/other 六个选项，选择后携带 expectedCategory 回传给后端记录纠错。详情卡片在用户点击“查看详情”时由后端返回的正文预览数据构建并新发，包含主题、收发件人、时间、分类与重要性、摘要及正文预览。")

h2("3.4 卡片刷新交互闭环设计")
para("飞书 card.action.trigger 回调有约 3 秒超时，超时会在卡片上显示红色错误；而后端真实业务（IMAP/数据库）可能慢于 3 秒。为兼顾“不超时”与“操作后卡片状态正确刷新”，本人设计了“先 toast、再异步刷新”的交互模式，其逻辑如下。")
code_block(r"""
handleCardAction(event):
  if action == feedback_wrong 且未带 expectedCategory:
      return { card: 分类选择卡片 }          # 唯一同步分支，不依赖后端
  scheduleBackendAndRefresh(...)              # 不 await，立即返回
  return { toast: "处理中..." }               # ① 毫秒级返回，让飞书停止倒计时

doBackendAndRefresh(...):                     # ② 后台执行
  result = forwardToBackend(payload+X-Bot-Secret)
  if not result.success: sendText(失败提示); return   # 失败可见化
  if action == view_detail:
      sendCard(详情卡片)                       # 新发，保留原卡片
  else:
      updateCard(messageId, buildEmailCard(result.email))  # ③ patch 原卡片刷新
      sendText(带主题的简短确认)
""")
caption("图 3-3　“先 toast 再异步刷新”交互闭环伪代码")
para("该设计同时解决了三个问题：一是规避了飞书 3 秒超时红色错误；二是保证所有按钮操作后卡片状态都能正确刷新（依赖从事件中稳健提取的 open_message_id）；三是将后端或刷新链路任意环节的失败以文字消息明确反馈给用户，不再静默吞错。这一模式是早期“为规避超时直接返回空响应导致卡片永不刷新”问题的根治方案。")

h2("3.5 大模型 Prompt 设计与参数调优")
para("分析效果直接取决于提交给大模型的提示词与模型参数。本人负责 OpenClaw + DeepSeek 分析的 Prompt 设计与参数调优（OpenClaw 子进程的工程接入由后端同学完成），目标是让模型在分类、重要性、摘要三项任务上稳定输出可解析、可用的结果。")
h3("3.5.1 Prompt 设计")
para("提示词采用强约束的结构化设计，要点包括：明确角色与任务；强制只返回单个合法 JSON 对象、禁止 Markdown 代码块与围栏、禁止解释性文字；要求对字符串中的双引号与换行正确转义以保证可被直接解析；给出精确的字段 schema；附上结构化的邮件信息并对正文做长度截断以控制 token 成本。其结构示意如下。")
code_block(r"""
请作为 EmailClaw 邮件分析 Agent 分析下面这封邮件。
你必须只返回一个合法的 JSON 对象，不要返回 Markdown 代码块/围栏或任何解释文字。
字符串中若含双引号或换行必须正确转义，确保整段可被 JSON.parse 解析。
字段格式：
{"category":"work|personal|shopping|marketing|spam|other",
 "confidence":0.8,"classificationReasoning":"...",
 "importance":7,"importanceReasoning":"...","summary":"..."}

UID/Message-ID/Subject/From/To/Date/Attachments: ...
Body: <正文截断至 6000 字符>
""")
caption("图 3-4　结构化分析提示词设计")
h3("3.5.2 参数与模型选型调优")
add_table(
    ["项", "取值/策略", "调优考量"],
    [
        ["模型", "deepseek/deepseek-chat", "最便宜的非推理模型，中文分类/摘要效果已足够，控成本"],
        ["正文截断", "前 6000 字符", "兼顾上下文充分性与 token 成本、超时风险"],
        ["超时", "OPENCLAW_TIMEOUT_MS=60000", "由默认 30s 上调至 60s，降低长邮件分析超时概率"],
        ["输出归一化", "类别越界归 other、置信度[0,1]、重要性[0,10]、缺省填默认", "保证下游拿到结构稳定的结果"],
    ],
    widths=[65, 180, 265], size=9,
)
caption("表 3-3　模型参数与选型调优")
h3("3.5.3 效果评估指标")
para("本人以可量化的现象作为调优是否生效的判定标准：分类结果不应清一色退化为“other”、重要性不应清一色退化为默认值“3”（这是提示词约束不足或解析失败时的典型表现）；模型输出应能被稳定解析为合法 JSON，不夹带 Markdown 围栏；分类理由与摘要应贴合邮件内容。配合后端的容错重试，最终有效解析率被拉到接近 100%；当大模型不可用时，系统回退本地关键词规则，仍能产出基础分类，保证体验不中断。")

# ====================================================================
# 第4章 系统实现
# ====================================================================
page_break()
h1("第4章 系统实现")
h2("4.1 引言")
para("本章介绍本人负责模块的实现情况。飞书机器人服务作为独立子项目 email_claw_bot 开发，基于 Node.js + 飞书 Node SDK，关键脚本通过语法检查。源码结构如下。")
code_block(r"""
email_claw_bot/src/
├── server.js                飞书 bot 入口(:3001)：HTTP 接口 + WSClient 事件分发
├── cards/emailCard.js       三类卡片构建器：邮件卡片/详情卡片/分类选择卡片
├── handlers/cardHandler.js  卡片按钮事件处理：先 toast 再异步 updateCard、转发后端
└── services/feishuClient.js 飞书 OpenAPI 封装：sendCard/updateCard/sendText/getMessage
""")
caption("图 4-1　飞书机器人服务源码结构")

h2("4.2 关键实现与运行")
h3("事件接收与卡片发送")
para("server.js 初始化飞书客户端并启动 WSClient 长连接，注册三类事件：用户进入单聊（发送欢迎语）、接收用户消息（回复提示）、以及核心的 card.action.trigger（交给 cardHandler 处理）。同时提供 HTTP 接口供后端调用。")
code_block(r"""
# 后端推送 → bot 发送卡片
POST :3001/api/notify-email   Body: { emailId, subject, from, category,
                                       importance, summary, isImportant, openId, ... }
# 无后端独立调试卡片样式
POST :3001/api/mock-card      Body: {}  (内置一封 mock 邮件数据)
# 健康检查
GET  :3001/ping               → { status: "ok", service: "EmailClaw Bot" }
""")
caption("图 4-2　飞书机器人 HTTP 接口")
h3("OpenClaw + DeepSeek 配置（模型侧）")
para("在模型侧，本人完成 DeepSeek API Key 配置、默认模型选定与 email-claw agent 创建，并单独自测连通性与分析 JSON 输出后再接入后端，关键命令如下。")
code_block(r"""
openclaw config set models.providers.deepseek.apiKey "sk-..."
openclaw config set agents.defaults.model "deepseek/deepseek-chat"
openclaw agents add email-claw --model deepseek/deepseek-chat --non-interactive
# 自测分析 JSON：openclaw agent --agent email-claw --json --local --message "..."
""")
caption("图 4-3　模型侧配置与自测")

h2("4.3 测试计划")
para("本人针对负责模块设计了卡片交互测试与模型效果评估。测试环境：Node 20 + 飞书企业自建应用；模型用 OpenClaw + DeepSeek。主要用例如下。")
add_table(
    ["编号", "测试项", "预期结果"],
    [
        ["TC-01", "卡片渲染（mock）", "mock 卡片样式、颜色、星级、按钮正确"],
        ["TC-02", "重要邮件高亮", "重要邮件红色卡头 + 🔥 前缀"],
        ["TC-03", "8 个按钮交互", "回传 action+emailId，先 toast 不超时"],
        ["TC-04", "卡片状态刷新", "操作后原卡片 updateCard 正确刷新"],
        ["TC-05", "查看详情", "新发详情卡片，原卡片按钮保留"],
        ["TC-06", "分类错误纠错", "弹分类选择卡片，选择后回传 expectedCategory"],
        ["TC-07", "失败可见化", "后端失败时返回明确文字提示"],
        ["TC-08", "模型效果评估", "分类非清一色 other、重要性非清一色 3、JSON 可解析"],
    ],
    widths=[45, 150, 315], size=9,
)
caption("表 4-1　测试用例")

h2("4.4 测试报告")
para("本人对负责模块进行了真实联调与评估，结果如下。")
bullet("卡片联调：三类卡片均能正确渲染；重要邮件红色高亮生效；八个按钮均能回传事件，先 toast 模式下不再出现飞书 3 秒红色超时；标记/归档/删除/纠错等操作后原卡片状态正确刷新；查看详情能新发详情卡片并保留原卡片按钮；分类错误能弹出选择卡片并回传正确分类；后端失败时用户能收到明确文字提示。")
bullet("模型效果评估：完成 Prompt 与参数调优后，分类结果不再清一色退化为 other、重要性不再清一色为 3，输出能稳定解析为合法 JSON、无 Markdown 围栏；分类理由与摘要贴合邮件内容；将大模型关闭后系统回退本地规则仍能正常产出基础分类。")
para("联调中本人定位并根治了“卡片不刷新”的前端缺陷，其跟踪情况如下。")
add_table(
    ["缺陷", "现象", "修复"],
    [
        ["卡片不刷新", "查看详情无响应、各操作后卡片状态不变、只弹默认 toast",
         "bot 端改为“先 toast 再异步 updateCard”：毫秒级返回 toast 停止倒计时，后台 await 后端结果后 patch 原卡片；详情改为新发卡片；失败以文字提示可见化"],
    ],
    widths=[55, 215, 240], size=9,
)
caption("表 4-2　缺陷跟踪与修复")
para("综合结果，本人负责的飞书前端已稳定打通“接收推送 → 展示卡片 → 用户点击 → 回传后端 → 刷新卡片/展示详情”的交互闭环，模型调优也使分析效果达到可用水平，共同支撑了 FR-2、FR-3、FR-4、FR-5 的用户侧呈现与分析质量。")

# ====================================================================
# 第5章 总结
# ====================================================================
page_break()
h1("第5章 总结")
para("在本次综合项目实践中，本人负责 EmailClaw 智能邮件管家系统的飞书机器人前端服务与大模型参数调优，完成了面向用户的交互入口与决定分析效果的模型侧工作。具体成果包括：基于飞书 Node SDK 与 WebSocket 长连接搭建的机器人服务、邮件通知/详情/分类选择三类交互式卡片、“先 toast 再异步刷新”的卡片交互闭环，以及 OpenClaw + DeepSeek 的 Prompt 设计与参数调优。")
para("本人工作的难点与收获主要集中在两方面：其一是飞书卡片的交互闭环，飞书 3 秒回调超时与后端较慢的真实业务之间存在矛盾，最初直接返回空响应导致卡片永不刷新，最终通过“先 toast 让飞书停止倒计时、再后台异步 updateCard 刷新原卡片”的模式才得以兼顾不超时与状态刷新，这让我体会到在受平台限制的前端交互中“响应时序设计”的重要性；其二是大模型的 Prompt 与参数调优，大模型并不会天然输出规范结果，通过强约束提示词、字段规范、正文截断与输出归一化，并配合容错重试，才把分类退化与非法 JSON 压到很低，让我认识到 AI 应用中“把模型输出驯化为稳定可用结构”往往比调用本身更关键。")
para("不足之处在于：当前分类纠错虽已回传记录，但尚未形成喂回模型的反馈学习闭环；卡片暂未支持超长正文的分页展示；FR-7 每日摘要卡片也有待实现。后续可在反馈学习、富交互卡片与定时摘要推送方向继续完善。")
para("通过这部分工作，本人完整经历了从需求分析、交互设计、编码实现到测试联调与缺陷修复的全过程，在第三方开放平台对接、交互时序设计、大模型提示工程与效果评估等方面都得到了切实锻炼，对面向用户的 AI 应用如何做到既好用又准确有了更深入的理解。")

doc.save(OUTPUT)
print("个人报告已生成：", OUTPUT)
